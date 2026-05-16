"""
导出工具模块
支持 TXT / MD / HTML / DOCX 四种格式
功能：
- 小说单章导出
- 小说全书导出（含目录）
- 角色扮演对话导出
"""

import os
import re

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.novel_manager import NovelManager


# ========== 共用模板 ==========

HTML_STYLE = """
<style>
  * { scrollbar-width: thin; }
  body {
    font-family: -apple-system, "Segoe UI", "Noto Sans SC", "Microsoft YaHei", sans-serif;
    background: #1a1a2e;
    color: #e0e0e0;
    padding: 40px;
    max-width: 900px;
    margin: 0 auto;
    line-height: 1.8;
    font-size: 15px;
  }
  h1 { color: #569cd6; font-size: 1.8em; border-bottom: 2px solid rgba(86,156,214,0.3); padding-bottom: 10px; margin-top: 0; }
  h2 { color: #569cd6; font-size: 1.4em; margin-top: 2em; border-bottom: 1px solid rgba(86,156,214,0.15); padding-bottom: 6px; }
  h3 { color: #dcdcaa; font-size: 1.15em; margin-top: 1.5em; }
  hr { border: none; height: 1px; background: linear-gradient(90deg,transparent,rgba(86,156,214,0.3),transparent); margin: 30px 0; }
  .meta { color: #888; font-size: 13px; margin-bottom: 30px; }
  .message { margin: 16px 0; padding: 12px 16px; border-radius: 8px; }
  .user-msg { background: rgba(86,156,214,0.08); border-left: 3px solid #569cd6; }
  .assistant-msg { background: rgba(212,220,170,0.06); border-left: 3px solid #dcdcaa; }
  .system-msg { background: rgba(106,153,85,0.08); border-left: 3px solid #6a9955; color: #6a9955; font-size: 13px; }
  .role-label { font-weight: 600; margin-bottom: 4px; }
  pre { background: #0d0d1a; border-radius: 8px; padding: 14px 18px; overflow-x: auto; font-size: 13px; border: 1px solid rgba(255,255,255,0.08); }
  code { background: rgba(86,156,214,0.12); border-radius: 4px; padding: 2px 7px; font-size: 13px; color: #dcdcaa; }
  pre code { background: transparent; padding: 0; color: #d4d4d4; }
  blockquote { border-left: 3px solid #569cd6; margin: 10px 0; padding: 8px 18px; color: #b0c4de; background: rgba(86,156,214,0.05); border-radius: 0 6px 6px 0; }
  table { border-collapse: collapse; margin: 14px 0; width: 100%; }
  th, td { border: 1px solid rgba(255,255,255,0.1); padding: 8px 12px; text-align: left; }
  th { background: #0d0d1a; color: #569cd6; }
  .toc a { color: #569cd6; text-decoration: none; }
  .toc a:hover { text-decoration: underline; }
  .footer { color: #555; font-size: 12px; text-align: center; margin-top: 50px; border-top: 1px solid rgba(255,255,255,0.05); padding-top: 20px; }
</style>
"""

HTML_WRAPPER = """<!DOCTYPE html>
<html><head><meta charset="utf-8">{style}</head><body>
{content}
<div class="footer">由 DeepSeek 多功能聊天客户端生成</div>
</body></html>"""


# ========== 小说导出 ==========

def _read_chapter_content(novel_manager: NovelManager, title: str, chapter_num: int) -> str | None:
    """读取章节的活跃版本内容"""
    return novel_manager.read_active_chapter(title, chapter_num)


def _build_chapter_meta(novel_manager: NovelManager, title: str, chapter_num: int) -> dict:
    """收集章节元信息"""
    meta = novel_manager.load_meta(title)
    chapters = novel_manager.list_chapters(title)
    ch_info = next((c for c in chapters if c["num"] == chapter_num), None)
    return {
        "book_title": title,
        "chapter_num": chapter_num,
        "chapter_title": ch_info["title"] if ch_info else f"第{chapter_num}章",
        "author": meta.author,
        "total_chapters": meta.total_chapters,
        "protagonist": meta.protagonist_bio,
        "background": meta.background_story,
    }


def export_chapter(
    novel_manager: NovelManager,
    title: str,
    chapter_num: int,
    fmt: str = "txt",
    output_path: str | None = None,
) -> str:
    """
    导出小说单章

    Args:
        novel_manager: NovelManager 实例
        title: 小说名称
        chapter_num: 章节编号
        fmt: 导出格式 (txt/md/html/docx)
        output_path: 指定输出路径，None 则自动生成

    Returns:
        导出文件的绝对路径
    """
    content = _read_chapter_content(novel_manager, title, chapter_num)
    if not content:
        raise FileNotFoundError(f"小说「{title}」第{chapter_num}章内容为空或不存在")

    meta = _build_chapter_meta(novel_manager, title, chapter_num)
    chapter_title = meta["chapter_title"]

    if output_path is None:
        output_path = _default_chapter_path(title, chapter_num, chapter_title, fmt)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    if fmt == "txt":
        text = _chapter_to_txt(content, meta)
        _write_text(output_path, text)
    elif fmt == "md":
        text = _chapter_to_md(content, meta)
        _write_text(output_path, text)
    elif fmt == "html":
        text = _chapter_to_html(content, meta)
        _write_text(output_path, text)
    elif fmt == "docx":
        _chapter_to_docx(content, meta, output_path)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")

    return os.path.abspath(output_path)


def _chapter_to_txt(content: str, meta: dict) -> str:
    lines = [
        f"{meta['book_title']} · {meta['chapter_title']}",
        "=" * 40,
        "",
        content,
    ]
    return "\n".join(lines)


def _chapter_to_md(content: str, meta: dict) -> str:
    return (
        f"# {meta['book_title']} · {meta['chapter_title']}\n\n"
        f"{content}\n"
    )


def _chapter_to_html(content: str, meta: dict) -> str:
    import markdown as md_lib
    body = md_lib.markdown(content, extensions=["fenced_code", "codehilite", "nl2br"])
    ch_title = f"{meta['book_title']} · {meta['chapter_title']}"
    html = f"""
<h1>{_escape_html(ch_title)}</h1>
<div class="meta">
  小说：{_escape_html(meta['book_title'])} | 章节：{_escape_html(meta['chapter_title'])}
</div>
<hr>
{body}
"""
    return HTML_WRAPPER.format(style=HTML_STYLE, content=html)


def _chapter_to_docx(content: str, meta: dict, output_path: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    p = doc.add_heading(f"{meta['book_title']} · {meta['chapter_title']}", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 正文
    for para in content.split("\n"):
        if not para.strip():
            doc.add_paragraph("")
        else:
            doc.add_paragraph(para)

    doc.save(output_path)


# ========== 全书导出 ==========

def export_book(
    novel_manager: NovelManager,
    title: str,
    fmt: str = "txt",
    output_path: str | None = None,
) -> str:
    """
    导出整本小说（所有章节按顺序拼接）

    Returns:
        导出文件的绝对路径
    """
    chapters = novel_manager.list_chapters(title)
    if not chapters:
        raise ValueError(f"小说「{title}」没有任何章节")

    meta = novel_manager.load_meta(title)

    if output_path is None:
        safe = _safe_filename(title)
        output_path = os.path.join(
            _book_dir(title),
            f"{safe}_全集.{fmt}",
        )

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    if fmt == "txt":
        text = _book_to_txt(novel_manager, title, chapters, meta)
        _write_text(output_path, text)
    elif fmt == "md":
        text = _book_to_md(novel_manager, title, chapters, meta)
        _write_text(output_path, text)
    elif fmt == "html":
        text = _book_to_html(novel_manager, title, chapters, meta)
        _write_text(output_path, text)
    elif fmt == "docx":
        _book_to_docx(novel_manager, title, chapters, meta, output_path)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")

    return os.path.abspath(output_path)


def _book_to_txt(novel_manager, title, chapters, meta) -> str:
    lines = [
        f"{'=' * 50}",
        f"  {title}",
        f"  作者: {meta.author}",
        f"  共 {len(chapters)} 章",
        "=" * 50,
        "",
    ]
    if meta.protagonist_bio:
        lines += ["【主角简介】", meta.protagonist_bio, ""]
    if meta.background_story:
        lines += ["【世界观设定】", meta.background_story, ""]
    lines.append("")

    for ch in chapters:
        content = _read_chapter_content(novel_manager, title, ch["num"])
        if content is None:
            continue
        lines += [
            "-" * 40,
            f"  第{ch['num']}章 {ch['title']}",
            "-" * 40,
            "",
            content,
            "",
        ]

    return "\n".join(lines)


def _book_to_md(novel_manager, title, chapters, meta) -> str:
    parts = [f"# {title}\n\n作者：{meta.author}  |  共 {len(chapters)} 章\n\n"]
    if meta.protagonist_bio:
        parts.append(f"## 主角简介\n\n{meta.protagonist_bio}\n\n")
    if meta.background_story:
        parts.append(f"## 世界观设定\n\n{meta.background_story}\n\n")
    parts.append("---\n\n")

    for ch in chapters:
        content = _read_chapter_content(novel_manager, title, ch["num"])
        if content is None:
            continue
        parts.append(f"## 第{ch['num']}章 {ch['title']}\n\n{content}\n\n")

    return "".join(parts)


def _book_to_html(novel_manager, title, chapters, meta) -> str:
    import markdown as md_lib
    body_parts = [
        f"<h1>{_escape_html(title)}</h1>",
        f'<div class="meta">作者：{_escape_html(meta.author)} | 共 {len(chapters)} 章</div>',
    ]
    if meta.protagonist_bio:
        body_parts.append(f"<h2>主角简介</h2><p>{_escape_html(meta.protagonist_bio)}</p>")
    if meta.background_story:
        body_parts.append(f"<h2>世界观设定</h2><p>{_escape_html(meta.background_story)}</p>")

    # 目录
    toc_parts = ['<div class="toc"><h2>目录</h2><ul>']
    for ch in chapters:
        toc_parts.append(
            f'<li><a href="#ch{ch["num"]}">第{ch["num"]}章 {_escape_html(ch["title"])}</a></li>'
        )
    toc_parts.append("</ul></div><hr>")
    body_parts.append("".join(toc_parts))

    for ch in chapters:
        content = _read_chapter_content(novel_manager, title, ch["num"])
        if content is None:
            continue
        body = md_lib.markdown(content, extensions=["fenced_code", "codehilite", "nl2br"])
        body_parts.append(
            f'<h2 id="ch{ch["num"]}">第{ch["num"]}章 {_escape_html(ch["title"])}</h2>\n{body}\n<hr>'
        )

    html = "\n".join(body_parts)
    return HTML_WRAPPER.format(style=HTML_STYLE, content=html)


def _book_to_docx(novel_manager, title, chapters, meta, output_path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 封面页
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    p = doc.add_paragraph(f"作者：{meta.author}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"共 {len(chapters)} 章")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 设定
    if meta.protagonist_bio:
        doc.add_heading("主角简介", level=1)
        doc.add_paragraph(meta.protagonist_bio)
    if meta.background_story:
        doc.add_heading("世界观设定", level=1)
        doc.add_paragraph(meta.background_story)

    # 目录
    doc.add_heading("目录", level=1)
    for ch in chapters:
        doc.add_paragraph(f"第{ch['num']}章 {ch['title']}")
    doc.add_page_break()

    # 章节
    for ch in chapters:
        content = _read_chapter_content(novel_manager, title, ch["num"])
        if content is None:
            continue
        doc.add_heading(f"第{ch['num']}章 {ch['title']}", level=1)
        for para in content.split("\n"):
            if not para.strip():
                doc.add_paragraph("")
            else:
                doc.add_paragraph(para)
        doc.add_page_break()

    doc.save(output_path)


# ========== 对话导出 ==========

def export_conversation(
    conversation_manager,
    conversation_id: str,
    fmt: str = "txt",
    output_path: str | None = None,
) -> str:
    """
    导出一段对话历史

    Returns:
        导出文件的绝对路径
    """
    record = conversation_manager.load_conversation(conversation_id)
    if not record:
        raise FileNotFoundError(f"对话「{conversation_id}」不存在")

    title = record.get("title", "未命名对话")
    messages = record.get("messages", [])
    model = record.get("model", "")
    strategy = record.get("strategy", "")
    char_desc = record.get("character_description", "")
    story_bg = record.get("story_background", "")
    created = record.get("created_at", "")
    updated = record.get("updated_at", "")

    if output_path is None:
        safe = _safe_filename(title)
        output_path = os.path.join(
            conversation_manager._root_dir,
            f"{safe}_导出.{fmt}",
        )

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    if fmt == "txt":
        text = _conv_to_txt(title, messages, model, strategy, char_desc, story_bg, created, updated)
        _write_text(output_path, text)
    elif fmt == "md":
        text = _conv_to_md(title, messages, model, strategy, char_desc, story_bg, created, updated)
        _write_text(output_path, text)
    elif fmt == "html":
        text = _conv_to_html(title, messages, model, strategy, char_desc, story_bg, created, updated)
        _write_text(output_path, text)
    elif fmt == "docx":
        _conv_to_docx(title, messages, model, strategy, char_desc, story_bg, created, updated, output_path)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")

    return os.path.abspath(output_path)


def _conv_metadata_block(model, strategy, char_desc, story_bg, created, updated) -> str:
    parts = []
    if model:
        parts.append(f"模型: {model}")
    if strategy:
        parts.append(f"模式: {strategy}")
    if created:
        parts.append(f"创建时间: {created}")
    if updated:
        parts.append(f"更新时间: {updated}")
    if char_desc:
        parts.append(f"角色描述: {char_desc}")
    if story_bg:
        parts.append(f"故事背景: {story_bg}")
    return "\n".join(parts)


def _conv_to_txt(title, messages, model, strategy, char_desc, story_bg, created, updated) -> str:
    lines = [
        f"{'=' * 50}",
        f"  {title}",
        "=" * 50,
        _conv_metadata_block(model, strategy, char_desc, story_bg, created, updated),
        "=" * 50,
        "",
    ]
    role_labels = {"user": "🧑 你", "assistant": "🤖 助手", "system": "[系统提示]"}
    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        label = role_labels.get(role, role)
        lines += [f"{label}:", content, ""]
    return "\n".join(lines)


def _conv_to_md(title, messages, model, strategy, char_desc, story_bg, created, updated) -> str:
    parts = [f"# {title}\n\n"]
    meta_block = _conv_metadata_block(model, strategy, char_desc, story_bg, created, updated)
    if meta_block:
        parts.append(f"{meta_block}\n\n---\n\n")

    role_labels = {"user": "🧑 你", "assistant": "🤖 助手", "system": "**系统提示**"}
    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        label = role_labels.get(role, role)
        if role == "system":
            parts.append(f"> {content}\n\n")
        else:
            parts.append(f"### {label}\n\n{content}\n\n")
    return "".join(parts)


def _conv_to_html(title, messages, model, strategy, char_desc, story_bg, created, updated) -> str:
    import markdown as md_lib
    body_parts = [
        f"<h1>{_escape_html(title)}</h1>",
        '<div class="meta">',
    ]
    if model:
        body_parts.append(f"模型: {_escape_html(model)}<br>")
    if strategy:
        body_parts.append(f"模式: {_escape_html(strategy)}<br>")
    if created:
        body_parts.append(f"创建时间: {created}<br>")
    if updated:
        body_parts.append(f"更新时间: {updated}<br>")
    if char_desc:
        body_parts.append(f"角色描述: {_escape_html(char_desc)}<br>")
    if story_bg:
        body_parts.append(f"故事背景: {_escape_html(story_bg)}<br>")
    body_parts.append("</div><hr>")

    role_labels = {"user": "🧑 你", "assistant": "🤖 助手", "system": "系统提示"}
    css_classes = {"user": "user-msg", "assistant": "assistant-msg", "system": "system-msg"}
    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        label = role_labels.get(role, role)
        cls = css_classes.get(role, "")
        if role == "system":
            body_parts.append(
                f'<div class="message {cls}"><div class="role-label">[{label}]</div>'
                f'{_escape_html(content)}</div>'
            )
        elif role == "user":
            body_parts.append(
                f'<div class="message {cls}"><div class="role-label">{label}</div>'
                f'{_escape_html(content)}</div>'
            )
        elif role == "assistant":
            md_body = md_lib.markdown(content, extensions=["fenced_code", "codehilite", "nl2br"])
            body_parts.append(
                f'<div class="message {cls}"><div class="role-label">{label}</div>{md_body}</div>'
            )

    html = "\n".join(body_parts)
    return HTML_WRAPPER.format(style=HTML_STYLE, content=html)


def _conv_to_docx(title, messages, model, strategy, char_desc, story_bg, created, updated, output_path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    doc.add_heading(title, level=1)
    doc.add_paragraph("")
    if model:
        doc.add_paragraph(f"模型: {model}")
    if strategy:
        doc.add_paragraph(f"模式: {strategy}")
    if created:
        doc.add_paragraph(f"创建时间: {created}")
    if updated:
        doc.add_paragraph(f"更新时间: {updated}")
    if char_desc:
        doc.add_paragraph(f"角色描述: {char_desc}")
    if story_bg:
        doc.add_paragraph(f"故事背景: {story_bg}")
    doc.add_paragraph("")

    role_labels = {"user": "🧑 你", "assistant": "🤖 助手", "system": "[系统提示]"}
    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        label = role_labels.get(role, role)
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        if role == "system":
            run.font.color.rgb = RGBColor(106, 153, 85)
            p.add_run(content)
        else:
            # 分段
            for para in content.split("\n"):
                doc.add_paragraph(para)
        doc.add_paragraph("")

    doc.save(output_path)


# ========== 通用辅助 ==========

def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _safe_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "_", name).strip()


def _write_text(path: str, text: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def _default_chapter_path(book_title: str, chapter_num: int, chapter_title: str, fmt: str) -> str:
    safe_book = _safe_filename(book_title)
    safe_ch = _safe_filename(chapter_title)
    return os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        "bookshelf",
        safe_book,
        f"第{chapter_num}章_{safe_ch}.{fmt}",
    )


def _book_dir(title: str) -> str:
    """获取小说目录路径"""
    return os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        "bookshelf",
        _safe_filename(title),
    )


EXPORT_FORMATS = ["txt", "md", "html", "docx"]
FORMAT_LABELS = {"txt": "纯文本 (.txt)", "md": "Markdown (.md)", "html": "HTML (.html)", "docx": "Word (.docx)"}
