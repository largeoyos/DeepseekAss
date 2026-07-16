
from __future__ import annotations

import asyncio
import html
import json
import os
import queue
import re
import shutil
import time
import uuid
import zipfile
from dataclasses import asdict
from email.parser import BytesParser
from email.policy import default as email_policy
from pathlib import Path

from fastapi import Depends, FastAPI, Header, HTTPException, Query, Request
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from core.auth_manager import AuthManager
from core.character_book import CharacterProfile, character_book_to_dict, dict_to_character_book
from core.chat_domain import ScenePreset, SceneState, SenderProfile, TurnPolicy, filter_fields, change_set_from_dict, apply_memory_change_set, revert_memory_change_set
from core.novel_manager import NovelMeta
from core.settings_manager import DEFAULT_PRESETS
from core.world_bible import audit_world_bible_consistency, dict_to_world_bible, world_bible_to_dict
from core.world_bible_diff import diff_world_bibles, summarize_world_bible_diff
from utils.export import export_book, export_chapter, export_conversation
from utils.summarize import detect_sections, split_text_locally
from web.services import WebApiConfigError, WebAuthError, WebRuntime, WebSensitiveError, WebUserContext, auxiliary_generation_model, body_generation_model, generation_params, masked_api_config, task_event_to_sse

class LoginRequest(BaseModel):
    username: str
    password: str

class SensitiveConfirmRequest(BaseModel):
    password: str

class ApiConfigRequest(BaseModel):
    text: dict = Field(default_factory=dict)
    image: dict = Field(default_factory=dict)

class ModelRequest(BaseModel):
    model: str = Field(min_length=1, max_length=120)

class SettingsUpdateRequest(BaseModel):
    settings: dict = Field(default_factory=dict)

class AgentWebTestRequest(BaseModel):
    query: str = "DeepseekAss 搜索测试"
class PresetRequest(BaseModel):
    name: str = Field(min_length=1, max_length=80)
    preset: dict = Field(default_factory=dict)

class PresetCurrentRequest(BaseModel):
    name: str = Field(min_length=1, max_length=80)

class ThemeRequest(BaseModel):
    theme: str = "dark"

class PasswordChangeRequest(BaseModel):
    old_password: str
    new_password: str = Field(min_length=6)

class DataImportRequest(BaseModel):
    overwrite: bool = True

class BookCreateRequest(BaseModel):
    title: str = Field(min_length=1, max_length=120)

class BookRenameRequest(BaseModel):
    new_title: str = Field(min_length=1, max_length=120)

class MetaUpdateRequest(BaseModel):
    protagonist_bio: str = ""
    background_story: str = ""
    writing_demand: str = ""
    author_plan: str = ""
    genre: str = ""
    style_tone: str = ""
    xp_mode: bool = False

class GenerateRequest(BaseModel):
    chapter_title: str = ""
    plot: str = ""
    target_words: int = Field(default=3000, ge=500, le=30000)

class AgentChapterPlanRequest(GenerateRequest):
    requirement: str = ""
    manual_entity_ids: list[str] = Field(default_factory=list)

class ExportRequest(BaseModel):
    fmt: str = "txt"
    chapter_num: int | None = None

class WorldSaveRequest(BaseModel):
    world: dict = Field(default_factory=dict)

class ContextPoliciesRequest(BaseModel):
    policies: dict = Field(default_factory=dict)

class WorldEntityRequest(BaseModel):
    category: str
    data: dict = Field(default_factory=dict)
    index: int | None = None

class AgentAdvisorRequest(BaseModel):
    message: str
    manual_references: list[str] = Field(default_factory=list)
    fiction_context: bool = True

class AgentAdviceSaveRequest(BaseModel):
    run_id: str = ""
    text: str
    title: str = "写作构思"

class AgentSessionCreateRequest(BaseModel):
    agent_kind: str = "writing_advisor"
    title: str = ""

class AgentWorkbenchRunRequest(BaseModel):
    message: str
    manual_references: list[str] = Field(default_factory=list)

class AgentRunControlRequest(BaseModel):
    payload: dict = Field(default_factory=dict)
class ContinuationSegmentRequest(BaseModel):
    text: str

class ContinuationImportRequest(BaseModel):
    title: str
    sections: list[dict]

class ContinuationAgentSegmentRequest(BaseModel):
    text: str = ""
    title: str = ""
    use_agent: bool = True

class ContinuationAnalyzeRequest(BaseModel):
    title: str = Field(min_length=1, max_length=120)
    sections: list[dict] = Field(default_factory=list)
    source_text: str = ""
    xp_mode: bool = False

class ContinuationSuggestRequest(BaseModel):
    title: str = ""
    setting: str = ""
    plot: str = ""
    world_data: dict = Field(default_factory=dict)
    xp_mode: bool = False

class ContinuationGenerateRequest(BaseModel):
    title: str = Field(min_length=1, max_length=120)
    source_text: str = ""
    chapter_title: str = ""
    requirement: str = ""
    plot: str = ""
    setting: str = ""
    target_words: int = Field(default=3000, ge=100, le=30000)
    xp_mode: bool = False
    chapter_mode: bool = True

class MarkdownWriteRequest(BaseModel):
    path: str
    content: str = ""

class MarkdownFolderRequest(BaseModel):
    path: str

class MarkdownRenameRequest(BaseModel):
    path: str
    new_path: str

class MarkdownExportRequest(BaseModel):
    path: str = ""
    folder: bool = False

class RoleProfileRequest(BaseModel):
    profile: dict = Field(default_factory=dict)

class SenderProfileRequest(BaseModel):
    profile: dict = Field(default_factory=dict)

class ScenePresetRequest(BaseModel):
    preset: dict = Field(default_factory=dict)

class ChatControlRequest(BaseModel):
    state: dict = Field(default_factory=dict)

class CharacterBookRequest(BaseModel):
    book: dict = Field(default_factory=dict)

class ConversationSaveRequest(BaseModel):
    record: dict = Field(default_factory=dict)

class RoleChatRequest(BaseModel):
    title: str = "Role Chat"
    message: str
    character_ids: list[str] = Field(default_factory=list)
    conversation_id: str = ""
    chat_type: str = "private"
    sender_name: str = "You"
    sender_profile: str = ""
    sender_profile_id: str = ""
    scene_state: dict = Field(default_factory=dict)
    turn_policy: dict = Field(default_factory=dict)
    required_responder_ids: list[str] = Field(default_factory=list)
    reply_mode: str = "character"
    narrator_enabled: bool = False

class ConversationExportRequest(BaseModel):
    fmt: str = "txt"

class ConversationBranchRequest(BaseModel):
    message_id: str = ""
    title: str = ""

class ConversationMessageRequest(BaseModel):
    content: str = ""
    requirement: str = ""
    title: str = ""

class MemoryChangeEditRequest(BaseModel):
    changes: list[dict] = Field(default_factory=list)
    apply_now: bool = False
class ChapterActionRequest(BaseModel):
    node_id: str = ""
    chapter_num: int | None = None
    version: int | None = None
    requirement: str = ""

class ChapterVariantRequest(BaseModel):
    mode: str = Field(default="polish", pattern="^(polish|rewrite)$")
    requirement: str = ""
    target_words: int = Field(default=3000, ge=0, le=30000)

class ChapterTreeActivateRequest(BaseModel):
    tree_id: str

class ChapterContentSaveRequest(BaseModel):
    content: str
    title: str = ""
    activate: bool = True

class ChapterSummarySaveRequest(BaseModel):
    summary: str = ""

class NodeExportRequest(BaseModel):
    fmt: str = "txt"

class SnapshotRequest(BaseModel):
    message: str = ""

class AgentChapterGenerateRequest(BaseModel):
    plan_id: str
    candidate_id: str = ""

class AgentChapterPlanCancelRequest(BaseModel):
    plan_id: str

class AgentChapterPlanRevisionRequest(BaseModel):
    plan_id: str
    candidate_id: str = ""
    instruction: str = Field(min_length=1, max_length=6000)

class AgentPolishPlanRequest(BaseModel):
    node_id: str
    requirement: str = ""

class AgentPolishGenerateRequest(BaseModel):
    plan_id: str

class AgentExtraPlanRequest(BaseModel):
    extra_type: str
    start_node_id: str = ""
    end_node_id: str = ""
    reference_node_id: str = ""
    title: str = ""
    plot: str = ""
    requirement: str = ""
    target_words: int = Field(default=5000, ge=500, le=50000)
    manual_entity_ids: list[str] = Field(default_factory=list)

class AgentExtraGenerateRequest(BaseModel):
    plan_id: str

class WorldDetailAnalyzeRequest(BaseModel):
    text: str
    source_run_id: str = ""

class WorldScopeConfirmRequest(BaseModel):
    change_set_id: str
    operations: list[dict] = Field(default_factory=list)

class ChangeApprovalRequest(BaseModel):
    change_set_id: str
    operation_ids: list[str] | None = None

class WorldEntityDeleteRequest(BaseModel):
    category: str
    index: int

class WorldEntityStateRequest(BaseModel):
    category: str
    index: int
    field: str
    value: object = True

class WorldResolveRequest(BaseModel):
    query: str

class WorldLockSettingRequest(BaseModel):
    topic: str
    passage: str = ""

class WorldForeshadowingRequest(BaseModel):
    hint: str
    relates_to: str = ""
    status: str = "open"
    next_step: str = ""
    reveal_rule: str = ""

class WorldRetrievalPreviewRequest(BaseModel):
    query: str = ""
    token_budget: int = Field(default=4000, ge=256, le=20000)

class WorldDuplicateDecisionRequest(BaseModel):
    candidate_id: str

class WorldUndoMergeRequest(BaseModel):
    merge_id: str = ""

class WorldCharacterMergeRequest(BaseModel):
    target_name: str
    merge_names: list[str] = Field(default_factory=list)

class WorldLocationMergeRequest(BaseModel):
    target_name: str
    merge_names: list[str] = Field(default_factory=list)

def create_app(runtime: WebRuntime | None = None) -> FastAPI:
    runtime = runtime or WebRuntime()
    app = FastAPI(title="DeepseekAss Web", version="0.3.0")
    app.state.runtime = runtime

    def token_from_header(authorization: str = Header(default="")) -> str:
        if not authorization.lower().startswith("bearer "):
            raise HTTPException(status_code=401, detail="未登录")
        return authorization.split(" ", 1)[1].strip()

    def current_context(token: str = Depends(token_from_header)) -> WebUserContext:
        try:
            return runtime.context_from_token(token)
        except WebAuthError as exc:
            raise HTTPException(status_code=401, detail=str(exc)) from exc

    def require_sensitive(ctx: WebUserContext, ticket: str) -> None:
        try:
            runtime.require_sensitive(ctx.username, ticket)
        except WebSensitiveError as exc:
            raise HTTPException(status_code=403, detail=str(exc)) from exc
    def text_client_and_model(ctx: WebUserContext):
        api_config = ctx.require_text_api()
        client = runtime.client_factory(api_config)
        model = (api_config.get("text") or {}).get("model") or ctx.settings.get("last_model") or "deepseek-v4-flash"
        return api_config, client, model

    @app.post("/api/auth/login", tags=["auth"])
    def login(payload: LoginRequest):
        try:
            return runtime.login(payload.username, payload.password)
        except WebAuthError as exc:
            raise HTTPException(status_code=401, detail=str(exc)) from exc

    @app.post("/api/auth/logout", tags=["auth"])
    def logout(token: str = Depends(token_from_header)):
        runtime.logout(token)
        return {"ok": True}

    @app.post("/api/auth/confirm", tags=["auth"])
    def confirm_sensitive(payload: SensitiveConfirmRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            ticket = runtime.confirm_sensitive(ctx.username, payload.password)
            return {"sensitive_ticket": ticket, "expires_in": 600}
        except WebSensitiveError as exc:
            raise HTTPException(status_code=403, detail=str(exc)) from exc

    @app.get("/api/session", tags=["auth"])
    def session(ctx: WebUserContext = Depends(current_context)):
        api_config = ctx.load_api_config()
        return {"user": {"username": ctx.username}, "api_configured": bool((api_config.get("text") or {}).get("api_key")), "settings": ctx.settings, "api": masked_api_config(api_config)}

    @app.get("/api/settings", tags=["settings"])
    def get_settings(ctx: WebUserContext = Depends(current_context)):
        return {"settings": ctx.settings_manager.load(), "api": masked_api_config(ctx.load_api_config())}

    @app.put("/api/settings", tags=["settings"])
    def save_settings(payload: SettingsUpdateRequest, ctx: WebUserContext = Depends(current_context)):
        settings = ctx.settings_manager.load()
        settings.update(payload.settings or {})
        ctx.save_settings(settings)
        return {"settings": ctx.settings}


    @app.put("/api/settings/model", tags=["settings"])
    def save_current_model(payload: ModelRequest, ctx: WebUserContext = Depends(current_context)):
        model = payload.model.strip()
        if not model:
            raise HTTPException(status_code=400, detail="模型名称不能为空")
        current = ctx.load_api_config()
        text_api = current.setdefault("text", {})
        text_api["model"] = model
        ctx.save_api_config(current)
        settings = ctx.settings_manager.load()
        settings["last_model"] = model
        custom_models = list(settings.get("custom_models") or [])
        if model not in custom_models:
            custom_models.append(model)
        settings["custom_models"] = custom_models
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"model": model, "settings": ctx.settings, "api": masked_api_config(current)}
    @app.get("/api/settings/presets", tags=["settings"])
    def get_presets(ctx: WebUserContext = Depends(current_context)):
        settings = ctx.settings_manager.load()
        return {"presets": settings.get("presets") or {}, "current_preset": settings.get("current_preset") or "", "theme": settings.get("theme", "dark"), "default_names": list(DEFAULT_PRESETS.keys())}

    @app.put("/api/settings/presets/current", tags=["settings"])
    def set_current_preset(payload: PresetCurrentRequest, ctx: WebUserContext = Depends(current_context)):
        preset_name = payload.name.strip()
        settings = ctx.settings_manager.load()
        presets = settings.get("presets") or {}
        if preset_name not in presets:
            raise HTTPException(status_code=404, detail="预设不存在")
        settings["current_preset"] = preset_name
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"current_preset": preset_name, "preset": presets[preset_name]}
    @app.put("/api/settings/presets/{name}", tags=["settings"])
    def save_preset(name: str, payload: PresetRequest, ctx: WebUserContext = Depends(current_context)):
        preset_name = (payload.name or name).strip()
        if not preset_name:
            raise HTTPException(status_code=400, detail="预设名称不能为空")
        preset = normalize_preset(payload.preset)
        settings = ctx.settings_manager.load()
        presets = dict(settings.get("presets") or {})
        presets[preset_name] = preset
        settings["presets"] = presets
        settings["current_preset"] = preset_name
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"presets": ctx.settings_manager.load().get("presets") or {}, "current_preset": preset_name}

    @app.delete("/api/settings/presets/{name}", tags=["settings"])
    def delete_preset(name: str, ctx: WebUserContext = Depends(current_context)):
        if name in DEFAULT_PRESETS:
            raise HTTPException(status_code=400, detail="默认预设不能删除，可恢复默认值")
        settings = ctx.settings_manager.load()
        presets = dict(settings.get("presets") or {})
        if name not in presets:
            raise HTTPException(status_code=404, detail="预设不存在")
        presets.pop(name, None)
        settings["presets"] = presets
        if settings.get("current_preset") == name:
            settings["current_preset"] = next(iter(presets), "")
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"presets": ctx.settings_manager.load().get("presets") or {}, "current_preset": settings.get("current_preset") or ""}

    @app.post("/api/settings/presets/reset", tags=["settings"])
    def reset_presets(ctx: WebUserContext = Depends(current_context)):
        settings = ctx.settings_manager.reset_presets()
        ctx.reload_settings()
        return {"presets": settings.get("presets") or {}, "current_preset": settings.get("current_preset") or ""}

    @app.put("/api/settings/theme", tags=["settings"])
    def save_theme(payload: ThemeRequest, ctx: WebUserContext = Depends(current_context)):
        theme = "light" if payload.theme == "light" else "dark"
        settings = ctx.settings_manager.load()
        settings["theme"] = theme
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"theme": theme}

    @app.get("/api/settings/agent-embedding", tags=["settings"])
    def get_agent_embedding_settings(ctx: WebUserContext = Depends(current_context)):
        settings = ctx.settings_manager.load()
        keys = [
            "novel_generation_mode", "controlled_agent_enabled", "agent_skills_enabled",
            "agent_runtime_backend", "retrieval_backend", "retrieval_default_limit",
            "retrieval_keyword_weight", "retrieval_semantic_weight", "retrieval_min_score",
            "framework_auto_fallback", "embedding_base_url", "embedding_api_key", "embedding_model",
            "embedding_batch_size", "embedding_timeout_seconds", "embedding_max_retries",
            "agent_web_enabled", "agent_web_endpoint", "agent_web_method", "agent_web_api_key",
            "agent_web_auth_header", "agent_web_auth_prefix", "agent_web_query_field",
            "agent_web_results_path", "agent_web_title_field", "agent_web_url_field",
            "agent_web_snippet_field", "agent_web_max_results", "agent_web_timeout_seconds",
        ]
        data = {key: settings.get(key) for key in keys if key in settings}
        for secret in ("embedding_api_key", "agent_web_api_key"):
            value = str(data.get(secret) or "")
            data[f"{secret}_configured"] = bool(value.strip())
            data[secret] = (value[:4] + "..." + value[-4:]) if len(value) > 8 else ("***" if value else "")
        return {"settings": data}

    @app.put("/api/settings/agent-embedding", tags=["settings"])
    def save_agent_embedding_settings(payload: SettingsUpdateRequest, sensitive_ticket: str = Header(default="", alias="X-Sensitive-Ticket"), ctx: WebUserContext = Depends(current_context)):
        incoming = dict(payload.settings or {})
        secret_keys = {"embedding_api_key", "agent_web_api_key"}
        if any(key in incoming and incoming.get(key) not in {"", "***", None} for key in secret_keys):
            require_sensitive(ctx, sensitive_ticket)
        settings = ctx.settings_manager.load()
        for key in secret_keys:
            if incoming.get(key) in {"", "***", None}:
                incoming.pop(key, None)
        settings.update(incoming)
        settings["controlled_agent_enabled"] = settings.get("novel_generation_mode") == "agent"
        ctx.settings_manager.save(settings)
        ctx.reload_settings()
        return {"settings": settings}

    @app.post("/api/settings/agent-web/test", tags=["settings"])
    def test_agent_web_search(payload: AgentWebTestRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.web_search import WebSearchClient, WebSearchConfig, WebSearchError

        settings = ctx.settings_manager.load()
        config = WebSearchConfig.from_settings(settings)
        if not config.enabled:
            raise HTTPException(status_code=400, detail="尚未启用受控网页搜索")
        if not config.endpoint:
            raise HTTPException(status_code=400, detail="尚未配置搜索 Endpoint")
        query = (payload.query or "DeepseekAss 搜索测试").strip() or "DeepseekAss 搜索测试"
        try:
            result = WebSearchClient(config).search(query, max_results=config.max_results)
        except WebSearchError as exc:
            raise HTTPException(status_code=400, detail=f"搜索测试失败：{exc}") from exc
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"搜索测试失败：{exc}") from exc
        results = result.get("results", []) if isinstance(result, dict) else []
        return {"ok": True, "query": query, "count": len(results), "results": results[: config.max_results]}

    @app.post("/api/settings/embedding/test", tags=["settings"])
    def test_embedding(ctx: WebUserContext = Depends(current_context)):
        ctx.reload_settings()
        backend = ctx.novel_manager.retrieval_backend()
        backend_name = str(getattr(backend, "backend_name", "") or "")
        if backend_name != "hybrid":
            raise HTTPException(status_code=400, detail="请先在设置中选择 LlamaIndex 混合检索")
        embedder = getattr(backend, "_embedder", None)
        get_query_embedding = getattr(embedder, "get_query_embedding", None)
        if not callable(get_query_embedding):
            raise HTTPException(status_code=400, detail="当前检索后端未暴露 Embedding 测试接口")
        try:
            vector = get_query_embedding("小说语义检索测试")
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Embedding 测试失败：{exc}") from exc
        dimension = len(vector) if hasattr(vector, "__len__") else 0
        return {"ok": True, "backend": backend_name, "dimension": dimension}
    @app.post("/api/books/{title}/retrieval/rebuild", tags=["settings"])
    def rebuild_retrieval(title: str, ctx: WebUserContext = Depends(current_context)):
        def target(handle):
            handle.progress("重建检索索引", percent=10, stage="检索")
            ctx.reload_settings()
            report = ctx.novel_manager.retrieval_backend().rebuild(title)
            handle.progress("检索索引重建完成", percent=100, stage="完成", data={"result": report})
            return report
        return {"task_id": runtime.start_task(ctx.username, f"重建《{title}》检索索引", target, metadata={"kind": "retrieval_rebuild", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/retrieval/clear", tags=["settings"])
    def clear_retrieval(title: str, ctx: WebUserContext = Depends(current_context)):
        backend = ctx.novel_manager.retrieval_backend()
        clear = getattr(backend, "clear", None)
        if callable(clear):
            clear(title)
        return {"ok": True}
    @app.put("/api/settings/api", tags=["settings"])
    def save_api_config(payload: ApiConfigRequest, sensitive_ticket: str = Header(default="", alias="X-Sensitive-Ticket"), ctx: WebUserContext = Depends(current_context)):
        require_sensitive(ctx, sensitive_ticket)
        current = ctx.load_api_config()
        for section in ("text", "image"):
            update = dict(getattr(payload, section) or {})
            if update.get("api_key") in {"", "***"}:
                update.pop("api_key", None)
            current[section].update(update)
        image = current.get("image") or {}
        image_values = [str(image.get(key) or "").strip() for key in ("api_key", "base_url", "model")]
        if any(image_values) and not all(image_values):
            raise HTTPException(status_code=400, detail="图片 API 如需启用，调用地址、API Key 和模型名称必须全部填写")
        ctx.save_api_config(current)
        text_api = current.get("text") or {}
        text_model = str(text_api.get("model") or "").strip()
        if text_model:
            settings = ctx.settings_manager.load()
            settings["last_model"] = text_model
            custom_models = list(settings.get("custom_models") or [])
            if text_model not in custom_models:
                custom_models.append(text_model)
            settings["custom_models"] = custom_models
            ctx.settings_manager.save(settings)
            ctx.reload_settings()
        return {"api": masked_api_config(current)}

    @app.post("/api/settings/password", tags=["settings"])
    def change_password(payload: PasswordChangeRequest, token: str = Depends(token_from_header), ctx: WebUserContext = Depends(current_context)):
        try:
            if not password_strength_ok(payload.new_password):
                raise ValueError("新密码至少 6 位，并且同时包含字母和数字")
            new_key = AuthManager.change_password(ctx.username, payload.old_password, payload.new_password)
            runtime.tokens.revoke(token)
            new_token = runtime.tokens.issue(ctx.username, new_key)
            return {"ok": True, "token": new_token, "user": {"username": ctx.username}}
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.post("/api/settings/data/export", tags=["settings"])
    def export_user_data(sensitive_ticket: str = Header(default="", alias="X-Sensitive-Ticket"), ctx: WebUserContext = Depends(current_context)):
        require_sensitive(ctx, sensitive_ticket)
        runtime.cleanup_exports(ctx)
        out = os.path.join(ctx.export_root, f"{safe_name(ctx.username)}_data.zip")
        root = os.path.abspath(ctx.user_dir)
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as archive:
            for current, _dirs, files in os.walk(root):
                for filename in files:
                    full = os.path.join(current, filename)
                    if os.path.abspath(full) == os.path.abspath(out):
                        continue
                    archive.write(full, os.path.relpath(full, root).replace("\\", "/"))
        download = runtime.register_download(ctx.username, out, os.path.basename(out), "application/zip")
        return {"download": download}

    @app.post("/api/settings/data/import", tags=["settings"])
    async def import_user_data(request: Request, sensitive_ticket: str = Header(default="", alias="X-Sensitive-Ticket"), ctx: WebUserContext = Depends(current_context)):
        require_sensitive(ctx, sensitive_ticket)
        uploaded = await read_single_upload(request)
        if not uploaded or not uploaded.get("content"):
            raise HTTPException(status_code=400, detail="请上传用户数据 ZIP 包")
        root = os.path.abspath(ctx.user_dir)
        imported = 0
        skipped = 0
        try:
            import io
            with zipfile.ZipFile(io.BytesIO(uploaded["content"]), "r") as archive:
                for member in archive.infolist():
                    name = str(member.filename or "").replace("\\", "/").strip("/")
                    if not name or name.startswith("../") or "/../" in name or os.path.isabs(name):
                        skipped += 1
                        continue
                    target = os.path.abspath(os.path.join(root, name))
                    if target != root and not target.startswith(root + os.sep):
                        skipped += 1
                        continue
                    if member.is_dir():
                        os.makedirs(target, exist_ok=True)
                        continue
                    os.makedirs(os.path.dirname(target), exist_ok=True)
                    with archive.open(member, "r") as source, open(target, "wb") as dest:
                        shutil.copyfileobj(source, dest)
                    imported += 1
        except zipfile.BadZipFile as exc:
            raise HTTPException(status_code=400, detail="ZIP 数据包无效") from exc
        return {"ok": True, "imported": imported, "skipped": skipped, "filename": uploaded.get("filename", "")}

    @app.post("/api/settings/data/clear", tags=["settings"])
    def clear_user_data(sensitive_ticket: str = Header(default="", alias="X-Sensitive-Ticket"), ctx: WebUserContext = Depends(current_context)):
        require_sensitive(ctx, sensitive_ticket)
        for name in os.listdir(ctx.user_dir):
            path = os.path.join(ctx.user_dir, name)
            try:
                if os.path.isdir(path):
                    shutil.rmtree(path)
                else:
                    os.remove(path)
            except FileNotFoundError:
                pass
        for name in ("bookshelf", "conversations"):
            os.makedirs(os.path.join(ctx.user_dir, name), exist_ok=True)
        return {"ok": True}

    @app.post("/api/settings/test-connection", tags=["settings"])
    def test_connection(ctx: WebUserContext = Depends(current_context)):
        try:
            api_config = ctx.require_text_api()
            client = runtime.client_factory(api_config)
            model = (api_config.get("text") or {}).get("model") or "deepseek-v4-flash"
            response = client.chat.completions.create(model=model, messages=[{"role": "user", "content": "只回复 OK"}], max_tokens=16, temperature=0)
            return {"ok": True, "reply": response.choices[0].message.content}
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.get("/api/books", tags=["books"])
    def list_books(ctx: WebUserContext = Depends(current_context)):
        return {"books": [{"title": title} for title in ctx.novel_manager.list_books()]}

    @app.post("/api/books", tags=["books"])
    def create_book(payload: BookCreateRequest, ctx: WebUserContext = Depends(current_context)):
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="书名不能为空")
        path = ctx.novel_manager.create_book(title)
        return {"title": title, "path": path}

    @app.patch("/api/books/{title}", tags=["books"])
    def rename_book(title: str, payload: BookRenameRequest, ctx: WebUserContext = Depends(current_context)):
        new_title = payload.new_title.strip()
        if not new_title:
            raise HTTPException(status_code=400, detail="书名不能为空")
        books = ctx.novel_manager.list_books()
        if title not in books:
            raise HTTPException(status_code=404, detail="书籍不存在")
        if new_title != title and new_title in books:
            raise HTTPException(status_code=400, detail="同名书籍已存在")
        ok = ctx.novel_manager.rename_book(title, new_title)
        if not ok:
            raise HTTPException(status_code=400, detail="重命名失败")
        return {"title": new_title}

    @app.delete("/api/books/{title}", tags=["books"])
    def delete_book(title: str, ctx: WebUserContext = Depends(current_context)):
        ok = ctx.novel_manager.delete_book(title)
        if not ok:
            raise HTTPException(status_code=404, detail="书籍不存在")
        return {"ok": True}

    @app.get("/api/books/{title}/meta", tags=["books"])
    def get_meta(title: str, ctx: WebUserContext = Depends(current_context)):
        if title not in ctx.novel_manager.list_books():
            raise HTTPException(status_code=404, detail="书籍不存在")
        return {"meta": serialize_meta(ctx.novel_manager.load_meta(title))}

    @app.put("/api/books/{title}/meta", tags=["books"])
    def save_meta(title: str, payload: MetaUpdateRequest, ctx: WebUserContext = Depends(current_context)):
        if title not in ctx.novel_manager.list_books():
            raise HTTPException(status_code=404, detail="书籍不存在")
        meta = ctx.novel_manager.save_meta(title, **model_data(payload))
        return {"meta": serialize_meta(meta)}

    @app.post("/api/books/{title}/generate", tags=["generation"])
    def generate(title: str, payload: GenerateRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            task_id = runtime.start_generation(ctx, title=title, chapter_title=payload.chapter_title.strip(), plot=payload.plot, target_words=payload.target_words)
            return {"task_id": task_id}
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
        except RuntimeError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

    @app.get("/api/books/{title}/chapters", tags=["chapters"])
    def list_chapters(title: str, ctx: WebUserContext = Depends(current_context)):
        return {"chapters": [normalize_chapter(item) for item in ctx.novel_manager.list_chapters(title)]}

    @app.get("/api/books/{title}/chapters/{chapter_num}", tags=["chapters"])
    def read_chapter(title: str, chapter_num: int, ctx: WebUserContext = Depends(current_context)):
        content = ctx.novel_manager.read_active_chapter(title, chapter_num)
        if content is None:
            raise HTTPException(status_code=404, detail="章节不存在")
        chapters = [normalize_chapter(item) for item in ctx.novel_manager.list_chapters(title)]
        current = next((item for item in chapters if int(item.get("chapter_num", 0)) == chapter_num), {})
        return {"chapter": current, "content": content}

    @app.get("/api/books/{title}/chapter-tree", tags=["chapters"])
    def chapter_tree(title: str, ctx: WebUserContext = Depends(current_context)):
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        return {"trees": ctx.novel_manager.list_chapter_trees(title), "nodes": ctx.novel_manager.list_chapter_tree_nodes(title), "active_path": meta.active_path, "active_tree_id": meta.active_tree_id, "target": ctx.novel_manager.get_active_generation_target(title)}

    @app.get("/api/books/{title}/nodes/{node_id}", tags=["chapters"])
    def read_node(title: str, node_id: str, ctx: WebUserContext = Depends(current_context)):
        content = ctx.novel_manager.read_chapter_node(title, node_id)
        if content is None:
            raise HTTPException(status_code=404, detail="节点不存在")
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        return {"node": meta.chapter_nodes.get(node_id), "content": content}

    @app.post("/api/books/{title}/chapter-trees/{tree_id}/activate", tags=["chapters"])
    def activate_chapter_tree(title: str, tree_id: str, ctx: WebUserContext = Depends(current_context)):
        if not ctx.novel_manager.switch_active_tree(title, tree_id):
            raise HTTPException(status_code=404, detail="阅读树不存在")
        try:
            ctx.novel_manager.rebuild_plot_summary_from_tree(title)
        except Exception:
            pass
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        return {"ok": True, "trees": ctx.novel_manager.list_chapter_trees(title), "nodes": ctx.novel_manager.list_chapter_tree_nodes(title), "active_path": meta.active_path, "active_tree_id": meta.active_tree_id}

    @app.get("/api/books/{title}/nodes/{node_id}/path", tags=["chapters"])
    def node_path(title: str, node_id: str, ctx: WebUserContext = Depends(current_context)):
        path = ctx.novel_manager.get_path_to_node(title, node_id)
        if not path:
            raise HTTPException(status_code=404, detail="节点不存在")
        return {"nodes": path}

    @app.get("/api/books/{title}/nodes/{node_id}/record", tags=["chapters"])
    def node_generation_record(title: str, node_id: str, ctx: WebUserContext = Depends(current_context)):
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        if node_id not in meta.chapter_nodes:
            raise HTTPException(status_code=404, detail="节点不存在")
        return {"record": ctx.novel_manager.load_node_generation_record(title, node_id) or {}}

    @app.put("/api/books/{title}/nodes/{node_id}/summary", tags=["chapters"])
    def save_node_summary(title: str, node_id: str, payload: ChapterSummarySaveRequest, ctx: WebUserContext = Depends(current_context)):
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        if node_id not in meta.chapter_nodes:
            raise HTTPException(status_code=404, detail="节点不存在")
        ctx.novel_manager.set_node_summary(title, node_id, payload.summary)
        try:
            ctx.novel_manager.rebuild_plot_summary_from_tree(title)
        except Exception:
            pass
        return {"ok": True, "node": ctx.novel_manager.ensure_chapter_tree(title).chapter_nodes.get(node_id)}

    @app.put("/api/books/{title}/nodes/{node_id}/content", tags=["chapters"])
    def save_node_content(title: str, node_id: str, payload: ChapterContentSaveRequest, ctx: WebUserContext = Depends(current_context)):
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        node = meta.chapter_nodes.get(node_id)
        if not node or node.get("virtual"):
            raise HTTPException(status_code=404, detail="节点不存在")
        if node.get("storage_kind") == "extra_uuid":
            saved_node = ctx.novel_manager.update_extra_node_content(title, node_id, payload.content, payload.title)
            return {"ok": True, "node_id": node_id, "version": saved_node.get("version"), "node": saved_node}
        chapter_num = int(node.get("chapter_num") or 0)
        if chapter_num <= 0:
            raise HTTPException(status_code=400, detail="节点章节号无效")
        chapter_title = payload.title.strip() or str(node.get("title") or f"第{chapter_num}章")
        version = ctx.novel_manager.get_next_version(title, chapter_num)
        _path, saved_version = ctx.novel_manager.save_chapter_version(title, chapter_num, chapter_title, payload.content, version=version, parent_id=node.get("parent_id"))
        if payload.activate:
            ctx.novel_manager.set_active_version(title, chapter_num, saved_version)
            try:
                ctx.novel_manager.rebuild_plot_summary_from_tree(title)
            except Exception:
                pass
        saved_node_id = ctx.novel_manager._node_id(chapter_num, saved_version)
        saved_meta = ctx.novel_manager.ensure_chapter_tree(title)
        return {"ok": True, "node_id": saved_node_id, "version": saved_version, "node": saved_meta.chapter_nodes.get(saved_node_id)}

    @app.post("/api/books/{title}/nodes/{node_id}/variant", tags=["chapters"])
    def chapter_node_variant(title: str, node_id: str, payload: ChapterVariantRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)

        def target(handle):
            from core.app_services import ChapterGenerationService
            from utils.prompts import Prompts
            from utils.supervision import supervise_chapter

            meta = ctx.novel_manager.load_meta(title)
            tree_meta = ctx.novel_manager.ensure_chapter_tree(title)
            node = tree_meta.chapter_nodes.get(node_id) or {}
            if not node or node.get("virtual") or node.get("storage_kind") == "extra_uuid":
                raise RuntimeError("请选择有效的正文章节节点")
            content = ctx.novel_manager.read_chapter_node(title, node_id) or ""
            if not content.strip():
                raise RuntimeError("当前章节正文为空")
            chapter_num = int(node.get("chapter_num") or 0)
            if chapter_num <= 0:
                raise RuntimeError("节点章节号无效")
            chapter_title = str(node.get("title") or node.get("display_label") or f"第{chapter_num}章")
            requirement = (payload.requirement or "").strip()
            mode = payload.mode if payload.mode in {"polish", "rewrite"} else "polish"
            handle.progress("准备章节变体", percent=8, stage="准备上下文")
            if mode == "polish":
                prompt = (
                    "请基于以下章节全文进行润色，保留核心剧情，不要输出解释。\n\n"
                    f"【润色要求】\n{requirement or '提升文笔、节奏和细节表现'}\n\n【原章节】\n{content}"
                )
                operation_label = "经典润色"
                generation_mode = "classic-polish-web"
                plot = "章节润色"
            else:
                summary = ctx.novel_manager.load_smart_summary(
                    title,
                    client=client,
                    next_chapter_num=chapter_num,
                    model=model,
                    global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                )
                prompt = (
                    f"请重写第 {chapter_num} 章「{chapter_title}」，不要输出解释。\n\n"
                    f"【前情提要】\n{summary}\n\n【重写要求】\n{requirement or '按既有设定重新组织本章内容'}\n\n"
                    f"【旧版本参考】\n{content[:4000]}"
                )
                operation_label = "经典重写"
                generation_mode = "classic-rewrite-web"
                plot = "章节重写"
            params = generation_params(ctx.settings, api_config)
            params["model"] = model
            handle.progress(f"{operation_label}正文", percent=20, stage="生成正文")
            candidate = runtime._stream_completion(
                handle,
                client,
                [{"role": "system", "content": Prompts.NOVEL_CHAPTER_WRITING}, {"role": "user", "content": prompt}],
                params,
            )
            if not candidate.strip():
                raise RuntimeError("模型未返回章节正文")
            handle.progress("监督修补", percent=62, stage="监督修补")
            supervision_report = {"status": "not_available"}
            final_content = candidate
            try:
                def supervision_client(_kind):
                    return client
                final_content, report = supervise_chapter(
                    supervision_client,
                    chapter_content=candidate,
                    chapter_title=chapter_title,
                    chapter_outline=plot,
                    requirements=requirement,
                    continuity_context=prompt,
                    target_words=int(payload.target_words or 0),
                    model=model,
                    temperature=params["temperature"],
                    global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                    xp_mode=bool(meta.xp_mode),
                    max_repair_rounds=2,
                )
                supervision_report = report.to_dict()
            except Exception as exc:
                supervision_report = {"status": "warning", "error": str(exc)}
            handle.progress("保存章节版本", percent=72, stage="保存章节")
            app_service = ChapterGenerationService(ctx.novel_manager)
            _path, saved_version = app_service.persist_chapter(
                title=title,
                chapter_num=chapter_num,
                chapter_title=chapter_title,
                content=final_content,
                version=ctx.novel_manager.get_next_version(title, chapter_num),
                parent_id=node.get("parent_id"),
                prompt=prompt,
                model=model,
                temperature=params["temperature"],
                top_p=params["top_p"],
                max_tokens=params["max_tokens"],
                frequency_penalty=params["frequency_penalty"],
                supervision_report=supervision_report,
                requirement=requirement,
                plot=plot,
                generation_mode=generation_mode,
            )
            try:
                ctx.novel_manager.update_generation_record(
                    title,
                    chapter_num,
                    saved_version,
                    operation="chapter_polish" if mode == "polish" else "chapter_rewrite",
                    polish_requirement=requirement if mode == "polish" else "",
                )
            except Exception:
                pass
            warnings = []
            handle.progress("生成章节摘要", percent=80, stage="生成摘要")
            try:
                ctx.novel_manager.generate_summary(
                    client,
                    final_content,
                    chapter_num,
                    chapter_title,
                    model=model,
                    global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                    xp_mode=bool(meta.xp_mode),
                    raise_on_error=True,
                )
            except Exception as exc:
                warnings.append(f"摘要生成失败：{exc}")
            handle.progress("更新世界书", percent=88, stage="更新世界书")
            try:
                app_service.world_bible.sync_chapter(
                    client,
                    title,
                    chapter_num,
                    saved_version,
                    final_content,
                    model=model,
                    global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                    xp_mode=bool(meta.xp_mode),
                )
            except Exception as exc:
                warnings.append(f"世界书更新失败：{exc}")
            handle.progress("创建项目快照", percent=95, stage="创建快照")
            snapshot_id = ""
            try:
                snapshot_id = app_service.create_auto_snapshot(title, chapter_num, saved_version).snapshot_id
            except Exception as exc:
                warnings.append(f"快照创建失败：{exc}")
            result = {
                "title": title,
                "node_id": ctx.novel_manager._node_id(chapter_num, saved_version),
                "chapter_num": chapter_num,
                "chapter_title": chapter_title,
                "version": saved_version,
                "mode": mode,
                "snapshot_id": snapshot_id,
                "warnings": warnings,
                "supervision_report": supervision_report,
                "preview": final_content[:240],
            }
            handle.progress(f"{operation_label}完成", percent=100, stage="完成", data={"result": result})
            return result

        label = "经典润色" if payload.mode == "polish" else "经典重写"
        return {"task_id": runtime.start_task(ctx.username, f"{label}《{title}》章节", target, metadata={"kind": f"chapter_{payload.mode}", "book": title, "node_id": node_id}, retryable=True)}
    @app.post("/api/books/{title}/nodes/{node_id}/export", tags=["export"])
    def export_node_task(title: str, node_id: str, payload: NodeExportRequest, ctx: WebUserContext = Depends(current_context)):
        fmt = normalize_fmt(payload.fmt)
        def target(handle):
            runtime.cleanup_exports(ctx)
            meta = ctx.novel_manager.ensure_chapter_tree(title)
            node = meta.chapter_nodes.get(node_id)
            content = ctx.novel_manager.read_chapter_node(title, node_id)
            if not node or content is None:
                raise FileNotFoundError("节点不存在")
            handle.progress("准备导出节点", percent=20, stage="导出")
            base = f"{safe_name(title)}_{safe_name(str(node.get('display_label') or node.get('title') or node_id))}.{fmt}"
            path = os.path.join(ctx.export_root, base)
            os.makedirs(ctx.export_root, exist_ok=True)
            write_node_export(path, fmt, title, node, content)
            download = runtime.register_download(ctx.username, path, os.path.basename(path), media_type_for(path))
            handle.progress("节点导出完成", percent=100, stage="完成", data={"download": download})
            return download
        return {"task_id": runtime.start_task(ctx.username, f"导出《{title}》节点", target, metadata={"kind": "node_export", "book": title, "node_id": node_id}, retryable=True)}

    @app.post("/api/books/{title}/nodes/{node_id}/activate", tags=["chapters"])
    def activate_node(title: str, node_id: str, ctx: WebUserContext = Depends(current_context)):
        if not ctx.novel_manager.switch_active_node(title, node_id):
            raise HTTPException(status_code=404, detail="节点不存在")
        try:
            ctx.novel_manager.rebuild_plot_summary_from_tree(title)
        except Exception:
            pass
        return {"ok": True}

    @app.delete("/api/books/{title}/nodes/{node_id}", tags=["chapters"])
    def delete_node(title: str, node_id: str, ctx: WebUserContext = Depends(current_context)):
        if not ctx.novel_manager.delete_chapter_node(title, node_id):
            raise HTTPException(status_code=404, detail="节点不存在或无法删除")
        return {"ok": True}

    @app.delete("/api/books/{title}/chapters/{chapter_num}", tags=["chapters"])
    def delete_chapter(title: str, chapter_num: int, ctx: WebUserContext = Depends(current_context)):
        if not ctx.novel_manager.delete_chapter(title, chapter_num):
            raise HTTPException(status_code=404, detail="章节不存在")
        return {"ok": True}

    @app.get("/api/books/{title}/active-path", tags=["chapters"])
    def active_path(title: str, ctx: WebUserContext = Depends(current_context)):
        return {"nodes": ctx.novel_manager.get_active_path_nodes(title)}

    @app.get("/api/books/{title}/chapters/{chapter_num}/versions", tags=["chapters"])
    def chapter_versions(title: str, chapter_num: int, ctx: WebUserContext = Depends(current_context)):
        active = ctx.novel_manager.get_active_version(title, chapter_num)
        return {"active": active, "versions": ctx.novel_manager.get_chapter_versions(title, chapter_num)}

    @app.get("/api/books/{title}/chapters/{chapter_num}/versions/{version}", tags=["chapters"])
    def read_chapter_version(title: str, chapter_num: int, version: int, ctx: WebUserContext = Depends(current_context)):
        content = ctx.novel_manager.read_chapter_version(title, chapter_num, version)
        if content is None:
            raise HTTPException(status_code=404, detail="章节版本不存在")
        return {"chapter_num": chapter_num, "version": version, "content": content}

    @app.post("/api/books/{title}/chapters/{chapter_num}/versions/{version}/activate", tags=["chapters"])
    def activate_chapter_version(title: str, chapter_num: int, version: int, ctx: WebUserContext = Depends(current_context)):
        node_id = ctx.novel_manager._node_id(chapter_num, version)
        if not ctx.novel_manager.switch_active_node(title, node_id):
            raise HTTPException(status_code=404, detail="章节版本不存在")
        ctx.novel_manager.rebuild_plot_summary_from_tree(title)
        return {"ok": True, "node_id": node_id}

    @app.delete("/api/books/{title}/chapters/{chapter_num}/versions/{version}", tags=["chapters"])
    def delete_version(title: str, chapter_num: int, version: int, ctx: WebUserContext = Depends(current_context)):
        if not ctx.novel_manager.delete_chapter_version(title, chapter_num, version):
            raise HTTPException(status_code=404, detail="章节版本不存在")
        return {"ok": True}

    @app.post("/api/books/{title}/chapters/rebuild-summary", tags=["chapters"])
    def rebuild_summary(title: str, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)
        def target(handle):
            handle.progress("重建活跃路径摘要", percent=10, stage="摘要")
            ctx.novel_manager.rebuild_summary_from_active(client, title, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(ctx.novel_manager.load_meta(title).xp_mode))
            handle.progress("摘要重建完成", percent=100, stage="完成")
            return {"ok": True}
        return {"task_id": runtime.start_task(ctx.username, f"重建《{title}》摘要", target, metadata={"kind": "summary_rebuild", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/world/rebuild", tags=["world"])
    def rebuild_world(title: str, payload: ChapterActionRequest, ctx: WebUserContext = Depends(current_context)):
        client = None
        model = str(ctx.settings.get("last_model") or "deepseek-v4-flash")
        if payload.requirement in {"extract_missing", "force_extract"}:
            _api, client, model = text_client_and_model(ctx)
        def target(handle):
            handle.progress("重建活跃路径世界书", percent=10, stage="世界书")
            report = ctx.novel_manager.rebuild_world_bible_from_active(client, title, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(ctx.novel_manager.load_meta(title).xp_mode), force_extract=payload.requirement == "force_extract", extract_missing=payload.requirement == "extract_missing")
            handle.progress("世界书重建完成", percent=100, stage="完成", data={"result": report})
            return report
        return {"task_id": runtime.start_task(ctx.username, f"重建《{title}》世界书", target, metadata={"kind": "world_rebuild", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/world/extract-node", tags=["world"])
    def extract_world_node(title: str, payload: ChapterActionRequest, ctx: WebUserContext = Depends(current_context)):
        _api, client, model = text_client_and_model(ctx)
        node_id = payload.node_id.strip()
        if not node_id:
            raise HTTPException(status_code=400, detail="请选择章节树节点")
        def target(handle):
            handle.progress("重提节点世界书", percent=10, stage="世界书")
            meta = ctx.novel_manager.ensure_chapter_tree(title)
            node = meta.chapter_nodes.get(node_id) or {}
            if node.get("storage_kind") == "extra_uuid":
                report = ctx.novel_manager.extract_world_bible_for_extra_node(client, title, node_id, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(ctx.novel_manager.load_meta(title).xp_mode), rebuild_active=True)
            else:
                report = ctx.novel_manager.extract_world_bible_for_node(client, title, node_id, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(ctx.novel_manager.load_meta(title).xp_mode))
            handle.progress("节点世界书已更新", percent=100, stage="完成", data={"result": report})
            return report
        return {"task_id": runtime.start_task(ctx.username, f"重提《{title}》节点世界书", target, metadata={"kind": "world_extract_node", "book": title, "node_id": node_id}, retryable=True)}
    @app.get("/api/books/{title}/context-preview", tags=["chapters"])
    def context_preview(title: str, chapter_title: str = "", plot: str = "", ctx: WebUserContext = Depends(current_context)):
        target = ctx.novel_manager.get_active_generation_target(title)
        report = ctx.novel_manager.context_assembler().assemble_chapter(title, int(target.get("chapter_num") or 1), chapter_title or f"第{target.get('chapter_num') or 1}章", plot, global_prompt=str(ctx.settings.get("global_user_prompt") or ""))
        return {"preview": report.preview(), "content": report.render(), "sections": [asdict(item) for item in report.sections]}

    @app.post("/api/books/{title}/export", tags=["export"])
    def export_book_task(title: str, payload: ExportRequest, ctx: WebUserContext = Depends(current_context)):
        fmt = normalize_fmt(payload.fmt)
        def target(handle):
            runtime.cleanup_exports(ctx)
            handle.progress("准备导出", percent=10, stage="导出")
            if payload.chapter_num:
                out = os.path.join(ctx.export_root, f"{safe_name(title)}_第{payload.chapter_num}章.{fmt}")
                path = export_chapter(ctx.novel_manager, title, int(payload.chapter_num), fmt, out)
            else:
                out = os.path.join(ctx.export_root, f"{safe_name(title)}_全书.{fmt}")
                path = export_book(ctx.novel_manager, title, fmt, out)
            download = runtime.register_download(ctx.username, path, os.path.basename(path), media_type_for(path))
            handle.progress("导出完成", percent=100, stage="完成", data={"download": download})
            return download
        return {"task_id": runtime.start_task(ctx.username, f"导出《{title}》", target, metadata={"kind": "export", "book": title}, retryable=True)}

    @app.get("/api/books/{title}/context-policies", tags=["chapters"])
    def get_context_policies(title: str, ctx: WebUserContext = Depends(current_context)):
        if title not in ctx.novel_manager.list_books():
            raise HTTPException(status_code=404, detail="书籍不存在")
        from core.context_assembler import _world_entities
        workspace = ctx.novel_manager.get_workspace(title)
        policies = workspace.load_context_policies()
        bible = ctx.novel_manager.load_world_bible(title)
        entities = []
        for entity_id, kind, name, _item in _world_entities(bible):
            entities.append({
                "entity_id": entity_id,
                "kind": kind,
                "name": name,
                "policy": normalize_context_policy(policies.get(entity_id) or {}),
            })
        return {"policies": policies, "entities": entities}

    @app.put("/api/books/{title}/context-policies", tags=["chapters"])
    def save_context_policies(title: str, payload: ContextPoliciesRequest, ctx: WebUserContext = Depends(current_context)):
        if title not in ctx.novel_manager.list_books():
            raise HTTPException(status_code=404, detail="书籍不存在")
        policies = {
            str(entity_id): normalize_context_policy(policy)
            for entity_id, policy in (payload.policies or {}).items()
            if str(entity_id).strip()
        }
        ctx.novel_manager.get_workspace(title).save_context_policies(policies)
        ctx.novel_manager.mark_retrieval_dirty(title, [{"source_type": "world_bible", "source_id": "context_policies"}])
        return {"policies": policies}

    @app.get("/api/books/{title}/world", tags=["world"])
    def get_world(title: str, ctx: WebUserContext = Depends(current_context)):
        bible = ctx.novel_manager.load_world_bible(title)
        return {"world": world_bible_to_dict(bible), "warnings": getattr(bible, "consistency_warnings", [])}

    @app.put("/api/books/{title}/world", tags=["world"])
    def save_world(title: str, payload: WorldSaveRequest, ctx: WebUserContext = Depends(current_context)):
        before = ctx.novel_manager.load_world_bible(title)
        after = dict_to_world_bible(payload.world)
        diff = diff_world_bibles(before, after)
        ctx.novel_manager.save_world_bible(title, after, force=True)
        return {"ok": True, "diff": [asdict(item) for item in diff], "summary": summarize_world_bible_diff(diff)}

    @app.post("/api/books/{title}/world/entity", tags=["world"])
    def upsert_world_entity(title: str, payload: WorldEntityRequest, ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        bucket = data.setdefault(payload.category, [])
        if not isinstance(bucket, list):
            raise HTTPException(status_code=400, detail="该分类不是列表")
        if payload.index is None:
            bucket.append(dict(payload.data or {}))
        else:
            if payload.index < 0 or payload.index >= len(bucket):
                raise HTTPException(status_code=404, detail="条目不存在")
            bucket[payload.index] = dict(payload.data or {})
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"world": world_bible_to_dict(bible)}

    @app.get("/api/books/{title}/world/audit", tags=["world"])
    def world_audit(title: str, ctx: WebUserContext = Depends(current_context)):
        return {"warnings": audit_world_bible_consistency(ctx.novel_manager.load_world_bible(title))}

    @app.post("/api/books/{title}/snapshots", tags=["snapshots"])
    def create_snapshot(title: str, message: str = "", ctx: WebUserContext = Depends(current_context)):
        snap = ctx.novel_manager.snapshot_service(title).create(message or "Web 手动快照", source="web")
        return {"snapshot": asdict(snap)}

    @app.get("/api/books/{title}/snapshots", tags=["snapshots"])
    def list_snapshots(title: str, ctx: WebUserContext = Depends(current_context)):
        return {"snapshots": [asdict(item) for item in ctx.novel_manager.snapshot_service(title).list()]}

    @app.get("/api/books/{title}/snapshots/{snapshot_id}/status", tags=["snapshots"])
    def snapshot_status(title: str, snapshot_id: str, ctx: WebUserContext = Depends(current_context)):
        return {"changes": ctx.novel_manager.snapshot_service(title).status(snapshot_id)}

    @app.post("/api/books/{title}/snapshots/{snapshot_id}/restore", tags=["snapshots"])
    def restore_snapshot(title: str, snapshot_id: str, ctx: WebUserContext = Depends(current_context)):
        restored = ctx.novel_manager.snapshot_service(title).restore(snapshot_id)
        return {"snapshot": asdict(restored)}

    @app.delete("/api/books/{title}/snapshots/{snapshot_id}", tags=["snapshots"])
    def delete_snapshot(title: str, snapshot_id: str, ctx: WebUserContext = Depends(current_context)):
        return {"ok": ctx.novel_manager.snapshot_service(title).delete(snapshot_id)}

    @app.delete("/api/books/{title}/world/entity", tags=["world"])
    def delete_world_entity(title: str, payload: WorldEntityDeleteRequest, ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        bucket = data.get(payload.category)
        if not isinstance(bucket, list):
            raise HTTPException(status_code=400, detail="该分类不是列表")
        if payload.index < 0 or payload.index >= len(bucket):
            raise HTTPException(status_code=404, detail="条目不存在")
        removed = bucket.pop(payload.index)
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "removed": removed, "world": world_bible_to_dict(bible)}
    @app.post("/api/books/{title}/world/entity/state", tags=["world"])
    def update_world_entity_state(title: str, payload: WorldEntityStateRequest, ctx: WebUserContext = Depends(current_context)):
        if payload.field not in {"hidden", "locked", "status"}:
            raise HTTPException(status_code=400, detail="不支持的状态字段")
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        bucket = data.get(payload.category)
        if not isinstance(bucket, list):
            raise HTTPException(status_code=400, detail="该分类不是列表")
        if payload.index < 0 or payload.index >= len(bucket):
            raise HTTPException(status_code=404, detail="条目不存在")
        item = bucket[payload.index]
        if not isinstance(item, dict):
            raise HTTPException(status_code=400, detail="该条目不是结构化对象")
        item[payload.field] = payload.value
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "world": world_bible_to_dict(bible), "item": item}

    @app.post("/api/books/{title}/world/resolve", tags=["world"])
    def resolve_world_items(title: str, payload: WorldResolveRequest, ctx: WebUserContext = Depends(current_context)):
        query = payload.query.strip().lower()
        if not query:
            raise HTTPException(status_code=400, detail="请输入剧情线或伏笔关键词")
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        changed = 0
        for item in data.get("active_plot_threads") or []:
            haystack = " ".join(str(item.get(key, "")) for key in ("name", "description", "expected_payoff", "payoff_hint")).lower()
            if query in haystack:
                item["status"] = "resolved"
                changed += 1
        for item in data.get("global_foreshadowing") or []:
            haystack = " ".join(str(item.get(key, "")) for key in ("hint", "relates_to", "next_step", "reveal_rule")).lower()
            if query in haystack:
                item["status"] = "resolved"
                changed += 1
        if not changed:
            raise HTTPException(status_code=404, detail="没有匹配到剧情线或伏笔")
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "changed": changed, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/lock-setting", tags=["world"])
    def lock_world_setting(title: str, payload: WorldLockSettingRequest, ctx: WebUserContext = Depends(current_context)):
        topic = payload.topic.strip()
        if not topic:
            raise HTTPException(status_code=400, detail="请输入设定主题")
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        bucket = data.setdefault("key_worldbuilding_passages", [])
        if not isinstance(bucket, list):
            raise HTTPException(status_code=400, detail="关键设定区不是列表")
        matched = next((item for item in bucket if isinstance(item, dict) and str(item.get("topic") or "").strip() == topic), None)
        if matched is None:
            matched = {"topic": topic, "passage": payload.passage.strip(), "chapter": int(data.get("last_updated_chapter") or 0)}
            bucket.append(matched)
        elif payload.passage.strip():
            matched["passage"] = payload.passage.strip()
        matched["locked"] = True
        matched["hidden"] = False
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "setting": matched, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/hide-low-priority", tags=["world"])
    def hide_low_priority_world_items(title: str, ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        changed = 0
        for item in data.get("characters") or []:
            if item.get("importance") == "minor" and not item.get("hidden"):
                item["hidden"] = True
                changed += 1
        for item in data.get("active_plot_threads") or []:
            if item.get("importance") == "minor" and item.get("status") != "active" and not item.get("hidden"):
                item["hidden"] = True
                changed += 1
        for item in data.get("locations") or []:
            low_signal = not item.get("key_details") and not item.get("atmosphere") and str(item.get("significance") or "").lower() in {"", "minor", "次要"}
            if low_signal and not item.get("hidden"):
                item["hidden"] = True
                changed += 1
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "changed": changed, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/foreshadowing", tags=["world"])
    def add_world_foreshadowing(title: str, payload: WorldForeshadowingRequest, ctx: WebUserContext = Depends(current_context)):
        hint = payload.hint.strip()
        if not hint:
            raise HTTPException(status_code=400, detail="伏笔内容不能为空")
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        chapter = int(data.get("last_updated_chapter") or 0)
        item = {"hint": hint, "relates_to": payload.relates_to.strip(), "status": payload.status.strip() or "open", "introduced_chapter": chapter, "last_touched_chapter": chapter, "next_step": payload.next_step.strip(), "reveal_rule": payload.reveal_rule.strip(), "hidden": False}
        data.setdefault("global_foreshadowing", []).append(item)
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "foreshadowing": item, "world": world_bible_to_dict(bible)}

    @app.get("/api/books/{title}/world/source", tags=["world"])
    def world_source_chapter(title: str, chapter: int = Query(default=0, ge=0), ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        if not chapter:
            chapter = int(data.get("last_updated_chapter") or 0)
        def match(item: dict, *keys: str) -> bool:
            return any(int(item.get(key) or 0) == chapter for key in keys)
        groups = {
            "characters": [item for item in data.get("characters", []) if match(item, "source_chapter", "last_updated_chapter", "first_appearance")],
            "locations": [item for item in data.get("locations", []) if match(item, "source_chapter", "last_updated_chapter", "first_appearance")],
            "timeline": [item for item in data.get("timeline", []) if match(item, "chapter", "source_chapter")],
            "active_plot_threads": [item for item in data.get("active_plot_threads", []) if match(item, "source_chapter", "opened_chapter", "last_touched_chapter")],
            "key_worldbuilding_passages": [item for item in data.get("key_worldbuilding_passages", []) if match(item, "chapter", "source_chapter")],
            "global_foreshadowing": [item for item in data.get("global_foreshadowing", []) if match(item, "introduced_chapter", "last_touched_chapter", "source_chapter")],
        }
        return {"chapter": chapter, "groups": groups}

    @app.post("/api/books/{title}/world/retrieval-preview", tags=["world"])
    def world_retrieval_preview(title: str, payload: WorldRetrievalPreviewRequest, ctx: WebUserContext = Depends(current_context)):
        from core.world_bible import format_relevant_world_bible_for_prompt
        bible = ctx.novel_manager.load_world_bible(title)
        meta = ctx.novel_manager.ensure_chapter_tree(title)
        active_chapters = {
            int((meta.chapter_nodes.get(node_id) or {}).get("chapter_num") or 0)
            for node_id in meta.active_path
            if int((meta.chapter_nodes.get(node_id) or {}).get("chapter_num") or 0) > 0
        }
        target_chapter = max(active_chapters or {int(getattr(bible, "last_updated_chapter", 0) or 0)})
        content, diagnostics = format_relevant_world_bible_for_prompt(bible, payload.query, active_chapters=active_chapters, target_chapter=target_chapter, token_budget=payload.token_budget, return_diagnostics=True)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"content": content, "diagnostics": diagnostics}

    @app.get("/api/books/{title}/world/facts", tags=["world"])
    def world_fact_history(title: str, entity_id: str = "", ctx: WebUserContext = Depends(current_context)):
        bible = ctx.novel_manager.load_world_bible(title)
        facts = [asdict(item) for item in getattr(bible, "facts", []) if not entity_id or item.subject_id == entity_id]
        return {"facts": facts}

    @app.get("/api/books/{title}/world/duplicates", tags=["world"])
    def world_duplicates(title: str, ctx: WebUserContext = Depends(current_context)):
        bible = ctx.novel_manager.load_world_bible(title)
        pending = [item for item in getattr(bible, "duplicate_candidates", []) if item.get("status", "pending") == "pending"]
        return {"pending": pending, "merge_history": [asdict(item) for item in getattr(bible, "merge_history", [])]}

    @app.post("/api/books/{title}/world/duplicates/confirm", tags=["world"])
    def confirm_world_duplicate(title: str, payload: WorldDuplicateDecisionRequest, ctx: WebUserContext = Depends(current_context)):
        from core.world_bible import confirm_duplicate_candidate
        bible = ctx.novel_manager.load_world_bible(title)
        if not confirm_duplicate_candidate(bible, payload.candidate_id):
            raise HTTPException(status_code=400, detail="候选已变化，无法安全合并")
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/duplicates/reject", tags=["world"])
    def reject_world_duplicate(title: str, payload: WorldDuplicateDecisionRequest, ctx: WebUserContext = Depends(current_context)):
        bible = ctx.novel_manager.load_world_bible(title)
        candidate = next((item for item in getattr(bible, "duplicate_candidates", []) if item.get("id") == payload.candidate_id), None)
        if not candidate:
            raise HTTPException(status_code=404, detail="候选不存在")
        candidate["status"] = "rejected"
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "candidate": candidate}

    @app.post("/api/books/{title}/world/merge/undo", tags=["world"])
    def undo_world_merge(title: str, payload: WorldUndoMergeRequest, ctx: WebUserContext = Depends(current_context)):
        from core.world_bible import undo_entity_merge
        bible = ctx.novel_manager.load_world_bible(title)
        if not undo_entity_merge(bible, payload.merge_id):
            raise HTTPException(status_code=404, detail="没有可撤销的实体合并")
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/characters/merge", tags=["world"])
    def merge_world_characters(title: str, payload: WorldCharacterMergeRequest, ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        characters = data.get("characters") or []
        if not isinstance(characters, list):
            raise HTTPException(status_code=400, detail="角色区不是列表")
        def keys(item: dict) -> set[str]:
            return {str(item.get("id") or "").lower(), str(item.get("name") or "").lower(), *[str(v).lower() for v in item.get("aliases") or []]}
        target = payload.target_name.strip().lower()
        base_index = next((idx for idx, item in enumerate(characters) if target and target in keys(item)), -1)
        if base_index < 0:
            raise HTTPException(status_code=404, detail="找不到主角色")
        merge_keys = [name.strip().lower() for name in payload.merge_names if name.strip()]
        if not merge_keys:
            raise HTTPException(status_code=400, detail="请输入要合并的角色名或别名")
        merge_indices: list[int] = []
        for key in merge_keys:
            idx = next((i for i, item in enumerate(characters) if i != base_index and key in keys(item)), -1)
            if idx >= 0 and idx not in merge_indices:
                merge_indices.append(idx)
        if not merge_indices:
            raise HTTPException(status_code=404, detail="没有找到可合并的角色")
        base = characters[base_index]
        removed_names: list[str] = []
        def append_unique_list(name: str, values: list) -> None:
            bucket = base.setdefault(name, [])
            for value in values or []:
                if value and value not in bucket:
                    bucket.append(value)
        def append_text(name: str, value: str, limit: int = 1200) -> None:
            value = str(value or "").strip()
            if not value:
                return
            current = str(base.get(name) or "").strip()
            if not current:
                base[name] = value[:limit]
            elif value not in current:
                base[name] = (current + "\n" + value)[:limit]
        importance_rank = {"minor": 0, "normal": 1, "major": 2}
        for idx in merge_indices:
            other = characters[idx]
            removed_names.append(str(other.get("name") or ""))
            append_unique_list("aliases", [other.get("name"), *(other.get("aliases") or [])])
            for field_name in ("traits", "notes", "arc"):
                append_text(field_name, other.get(field_name, ""))
            for field_name in ("motivation", "current_location", "current_goal", "current_emotion", "recent_action", "knowledge_state"):
                if not base.get(field_name) and other.get(field_name):
                    base[field_name] = other.get(field_name)
            for field_name in ("key_details", "key_dialogues", "unresolved_conflicts"):
                append_unique_list(field_name, other.get(field_name) or [])
            if importance_rank.get(other.get("importance"), 1) > importance_rank.get(base.get("importance"), 1):
                base["importance"] = other.get("importance")
            if base.get("status", "alive") == "alive" and other.get("status") and other.get("status") != "alive":
                base["status"] = other.get("status")
            for field_name in ("first_appearance", "source_chapter"):
                other_value = int(other.get(field_name) or 0)
                base_value = int(base.get(field_name) or 0)
                if other_value and (not base_value or other_value < base_value):
                    base[field_name] = other_value
        removed_set = {name for name in removed_names if name}
        for character in characters:
            for rel in character.get("relationships") or []:
                if isinstance(rel, dict) and rel.get("target") in removed_set:
                    rel["target"] = base.get("name")
        data["characters"] = [item for idx, item in enumerate(characters) if idx not in set(merge_indices)]
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "merged": removed_names, "world": world_bible_to_dict(bible)}

    @app.post("/api/books/{title}/world/locations/merge", tags=["world"])
    def merge_world_locations(title: str, payload: WorldLocationMergeRequest, ctx: WebUserContext = Depends(current_context)):
        data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
        locations = data.get("locations") or []
        if not isinstance(locations, list):
            raise HTTPException(status_code=400, detail="地点区不是列表")
        def keys(item: dict) -> set[str]:
            return {str(item.get("id") or "").lower(), str(item.get("name") or "").lower()}
        target = payload.target_name.strip().lower()
        base_index = next((idx for idx, item in enumerate(locations) if target and target in keys(item)), -1)
        if base_index < 0:
            raise HTTPException(status_code=404, detail="找不到主地点")
        merge_keys = [name.strip().lower() for name in payload.merge_names if name.strip()]
        if not merge_keys:
            raise HTTPException(status_code=400, detail="请输入要合并的地点名或 ID")
        merge_indices: list[int] = []
        for key in merge_keys:
            idx = next((i for i, item in enumerate(locations) if i != base_index and key in keys(item)), -1)
            if idx >= 0 and idx not in merge_indices:
                merge_indices.append(idx)
        if not merge_indices:
            raise HTTPException(status_code=404, detail="没有找到可合并的地点")
        base = locations[base_index]
        removed_names: list[str] = []
        def append_unique_list(name: str, values: list) -> None:
            bucket = base.setdefault(name, [])
            for value in values or []:
                if value and value not in bucket:
                    bucket.append(value)
        def append_text(name: str, value: str, limit: int = 1200) -> None:
            value = str(value or "").strip()
            if not value:
                return
            current = str(base.get(name) or "").strip()
            if not current:
                base[name] = value[:limit]
            elif value not in current:
                base[name] = (current + "\n" + value)[:limit]
        for idx in merge_indices:
            other = locations[idx]
            removed_names.append(str(other.get("name") or ""))
            for field_name, limit in (("description", 1600), ("significance", 900), ("atmosphere", 700)):
                append_text(field_name, other.get(field_name, ""), limit)
            append_unique_list("key_details", other.get("key_details") or [])
            append_unique_list("source_refs", other.get("source_refs") or [])
            for field_name in ("first_appearance", "source_chapter"):
                other_value = int(other.get(field_name) or 0)
                base_value = int(base.get(field_name) or 0)
                if other_value and (not base_value or other_value < base_value):
                    base[field_name] = other_value
            for field_name in ("last_updated_chapter", "last_updated_version"):
                other_value = int(other.get(field_name) or 0)
                if other_value and other_value > int(base.get(field_name) or 0):
                    base[field_name] = other_value
            base["hidden"] = bool(base.get("hidden")) and bool(other.get("hidden"))
            base["locked"] = bool(base.get("locked")) or bool(other.get("locked"))
            try:
                base["confidence"] = max(float(base.get("confidence") or 0), float(other.get("confidence") or 0))
            except Exception:
                pass
            for fact in data.get("facts") or []:
                if isinstance(fact, dict) and fact.get("subject_id") == other.get("id"):
                    fact["subject_id"] = base.get("id")
        data["locations"] = [item for idx, item in enumerate(locations) if idx not in set(merge_indices)]
        bible = dict_to_world_bible(data)
        ctx.novel_manager.save_world_bible(title, bible, force=True)
        return {"ok": True, "merged": removed_names, "world": world_bible_to_dict(bible)}
    @app.post("/api/books/{title}/agent/advisor", tags=["agent"])
    def ask_advisor(title: str, payload: AgentAdvisorRequest, ctx: WebUserContext = Depends(current_context)):
        api_config = ctx.require_text_api()
        def target(handle):
            from core.agent.advisor import AdvisorRequest, WritingAdvisorService
            client = runtime.client_factory(api_config)
            model = (api_config.get("text") or {}).get("model") or "deepseek-v4-flash"
            handle.progress("Agent 顾问思考中", percent=20, stage="Agent")
            result = WritingAdvisorService(ctx.novel_manager, client, ctx.conversation_manager).ask(AdvisorRequest(title, payload.message, model, ctx.settings, payload.manual_references, payload.fiction_context))
            handle.progress("Agent 顾问完成", percent=100, stage="完成", data={"result": asdict(result)})
            return asdict(result)
        return {"task_id": runtime.start_task(ctx.username, f"Agent 顾问《{title}》", target, metadata={"kind": "agent_advisor", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/agent/chapter/plan", tags=["agent"])
    def agent_chapter_plan(title: str, payload: AgentChapterPlanRequest, ctx: WebUserContext = Depends(current_context)):
        api_config = ctx.require_text_api()
        def target(handle):
            from core.agent.chapter_generation import AgentChapterGenerationService, AgentChapterPlan, AgentChapterRequest
            client = runtime.client_factory(api_config)
            model = (api_config.get("text") or {}).get("model") or "deepseek-v4-flash"
            planning_model = auxiliary_generation_model(ctx.settings, model)
            target_info = ctx.novel_manager.get_active_generation_target(title)
            chapter_num = int(target_info.get("chapter_num") or 1)
            chapter_title = payload.chapter_title or f"第{chapter_num}章"
            handle.progress("Agent 规划章节", percent=15, stage="Agent 规划")
            plan = AgentChapterGenerationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)), multi_plan_enabled=bool(ctx.settings.get("agent_multi_plan_enabled", False))).prepare(AgentChapterRequest(title, chapter_num, chapter_title, payload.plot, payload.requirement, payload.target_words, planning_model, payload.manual_entity_ids, str(ctx.settings.get("global_user_prompt") or "")))
            plan_data = plan.to_dict()
            for candidate in plan_data.get("candidate_plans") or []:
                candidate["rendered"] = AgentChapterPlan.from_dict(candidate).render()
            handle.progress("Agent 规划完成", percent=100, stage="完成", data={"plan": plan_data, "rendered": plan.render()})
            return {"plan": plan_data, "rendered": plan.render()}
        return {"task_id": runtime.start_task(ctx.username, f"Agent 规划《{title}》", target, metadata={"kind": "agent_chapter_plan", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/agent/chapter/cancel", tags=["agent"])
    def cancel_agent_chapter_plan(title: str, payload: AgentChapterPlanCancelRequest, ctx: WebUserContext = Depends(current_context)):
        workspace = ctx.novel_manager.get_workspace(title)
        path = f"{workspace.agent_root}/chapter_runs/{payload.plan_id}.json"
        record = workspace.storage.read_json(path, default={}) or {}
        if not record:
            raise HTTPException(status_code=404, detail="Agent 章节计划不存在")
        record.update({"status": "cancelled", "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S")})
        workspace.storage.write_json(path, record)
        return {"ok": True, "plan_id": payload.plan_id, "status": "cancelled"}
    @app.post("/api/books/{title}/agent/chapter/revise", tags=["agent"])
    def agent_chapter_revise(title: str, payload: AgentChapterPlanRevisionRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)

        def target(handle):
            from core.agent.chapter_generation import AgentChapterGenerationService, AgentChapterPlan, AgentChapterRequest

            workspace = ctx.novel_manager.get_workspace(title)
            record = workspace.storage.read_json(
                f"{workspace.agent_root}/chapter_runs/{payload.plan_id}.json", default={}
            ) or {}
            if not record:
                raise RuntimeError("Agent 章节计划不存在")
            request_data = dict(record.get("request") or {})
            request_data["model"] = auxiliary_generation_model(ctx.settings, model)
            request = AgentChapterRequest(**request_data)
            stored_plan = AgentChapterPlan.from_dict(dict(record.get("plan") or {}))
            candidates = list(stored_plan.candidate_plans or [stored_plan.to_dict(include_candidates=False)])
            selected_id = payload.candidate_id or stored_plan.candidate_id
            selected = next(
                (item for item in candidates if str(item.get("candidate_id") or "") == selected_id),
                None,
            )
            if selected is None:
                raise RuntimeError("所选 Agent 章节方案不存在")
            current_plan = AgentChapterPlan.from_dict(selected)
            current_plan.candidate_plans = candidates
            current_plan.recommended_candidate_id = stored_plan.recommended_candidate_id
            handle.progress("Agent 正在修改章节计划", percent=20, stage="计划修改")
            revised = AgentChapterGenerationService(
                ctx.novel_manager,
                client,
                skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)),
                multi_plan_enabled=False,
            ).revise_plan(request, current_plan, payload.instruction)
            plan_data = revised.to_dict()
            for candidate in plan_data.get("candidate_plans") or []:
                candidate["rendered"] = AgentChapterPlan.from_dict(candidate).render()
            handle.progress(
                "Agent 章节计划已修改", percent=100, stage="完成",
                data={"plan": plan_data, "rendered": revised.render()},
            )
            return {"plan": plan_data, "rendered": revised.render()}

        return {"task_id": runtime.start_task(
            ctx.username,
            f"Agent 修改计划《{title}》",
            target,
            metadata={"kind": "agent_chapter_revise", "book": title, "plan_id": payload.plan_id},
            retryable=True,
        )}
    @app.post("/api/books/{title}/agent/sessions", tags=["agent"])
    def create_agent_session(title: str, payload: AgentSessionCreateRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.profiles import AGENT_PROFILES
        from core.agent.repository import AgentRepository
        kind = payload.agent_kind if payload.agent_kind in AGENT_PROFILES else "writing_advisor"
        manifest = ctx.novel_manager.ensure_workspace(title)
        repo = AgentRepository(ctx.novel_manager.get_workspace(title))
        session = repo.create_session(manifest.book_id, title, kind, payload.title or AGENT_PROFILES[kind].display_name)
        return {"session": asdict(session)}

    @app.get("/api/books/{title}/agent/sessions/{session_id}", tags=["agent"])
    def get_agent_session(title: str, session_id: str, ctx: WebUserContext = Depends(current_context)):
        from core.agent.repository import AgentRepository
        repo = AgentRepository(ctx.novel_manager.get_workspace(title))
        session = repo.load_session(session_id)
        if session is None:
            raise HTTPException(status_code=404, detail="Agent 会话不存在")
        return {"session": asdict(session)}

    @app.post("/api/books/{title}/agent/sessions/{session_id}/run", tags=["agent"])
    def run_agent_session(title: str, session_id: str, payload: AgentWorkbenchRunRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.backends import build_agent_backend
            from core.agent.domain_tools import build_domain_tool_registry
            from core.agent.repository import AgentRepository
            from core.agent.types import AgentRunRequest
            repo = AgentRepository(ctx.novel_manager.get_workspace(title))
            session = repo.load_session(session_id)
            if session is None:
                raise RuntimeError("Agent 会话不存在")
            manifest = ctx.novel_manager.ensure_workspace(title)
            active_run_id = ""
            backend_holder = {"backend": None}
            def event_sink(event):
                nonlocal active_run_id
                if not active_run_id:
                    active_run_id = event.run_id
                    runtime.register_agent_backend(ctx.username, event.run_id, backend_holder["backend"])
                percent = min(95, 10 + int(event.sequence) * 5)
                handle.progress(f"Agent {event.event_type}", percent=percent, stage=event.event_type, data={"agent_event": asdict(event)})
            backend, status = build_agent_backend(
                settings=ctx.settings,
                novel_manager=ctx.novel_manager,
                client=client,
                tool_registry=build_domain_tool_registry(ctx.novel_manager, ctx.conversation_manager),
                event_sink=event_sink,
                skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)),
            )
            backend_holder["backend"] = backend
            request = AgentRunRequest(manifest.book_id, session_id, session.agent_kind, payload.message, payload.manual_references, model=model, book_title=title)
            result_run = None
            try:
                run = backend.run(request)
                result_run = run
                if repo.load_run(run.run_id) is None:
                    repo.save_run(run)
                result = {"run": asdict(run), "backend": asdict(status)}
                handle.progress("Agent 运行完成", percent=100, stage=run.status, data={"result": result})
                return result
            finally:
                if active_run_id and (result_run is None or result_run.status != "waiting_approval"):
                    runtime.unregister_agent_backend(active_run_id)
        return {"task_id": runtime.start_task(ctx.username, f"Agent 工作台《{title}》", target, metadata={"kind": "agent_workbench_run", "book": title, "session_id": session_id}, retryable=True)}

    @app.get("/api/books/{title}/agent/runs/{run_id}", tags=["agent"])
    def get_agent_run(title: str, run_id: str, ctx: WebUserContext = Depends(current_context)):
        from core.agent.repository import AgentRepository
        repo = AgentRepository(ctx.novel_manager.get_workspace(title))
        run = repo.load_run(run_id)
        if run is None:
            raise HTTPException(status_code=404, detail="Agent 运行不存在")
        workspace = ctx.novel_manager.get_workspace(title)
        ledger = workspace.storage.read_json(f"{workspace.agent_root}/ledger/{run_id}.json", default={}) or {}
        return {"run": asdict(run), "events": ledger.get("events", []) if isinstance(ledger, dict) else []}

    @app.post("/api/books/{title}/agent/runs/{run_id}/{action}", tags=["agent"])
    def control_agent_run(title: str, run_id: str, action: str, payload: AgentRunControlRequest | None = None, ctx: WebUserContext = Depends(current_context)):
        if action not in {"pause", "resume", "cancel"}:
            raise HTTPException(status_code=400, detail="不支持的 Agent 控制动作")
        control_payload = payload.payload if payload else {}
        if not runtime.control_agent_run(ctx.username, run_id, action, control_payload):
            raise HTTPException(status_code=404, detail="Agent 运行不在活动状态")
        return {"ok": True, "action": action}
    @app.post("/api/books/{title}/agent/advice", tags=["agent"])
    def save_agent_advice(title: str, payload: AgentAdviceSaveRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.advisor import WritingAdvisorService
        artifact_id = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager).save_advice(title, payload.run_id, payload.text, payload.title or "写作构思")
        advice = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager).list_advice(title)
        return {"artifact_id": artifact_id, "advice": advice}

    @app.delete("/api/books/{title}/agent/advisor/history/{message_index}", tags=["agent"])
    def delete_agent_advisor_history_message(title: str, message_index: int, ctx: WebUserContext = Depends(current_context)):
        from core.agent.advisor import WritingAdvisorService
        service = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager)
        if not service.delete_history_message(title, message_index):
            raise HTTPException(status_code=404, detail="顾问消息不存在")
        return {"ok": True, "advisor_history": service.list_history(title)}

    @app.delete("/api/books/{title}/agent/advisor/history", tags=["agent"])
    def clear_agent_advisor_history(title: str, ctx: WebUserContext = Depends(current_context)):
        from core.agent.advisor import WritingAdvisorService
        service = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager)
        removed = service.clear_history(title)
        return {"ok": True, "removed": removed, "advisor_history": service.list_history(title)}
    @app.get("/api/books/{title}/agent/state", tags=["agent"])
    def agent_state(title: str, ctx: WebUserContext = Depends(current_context)):
        from core.agent.profiles import AGENT_PROFILES
        from core.agent.repository import AgentRepository
        from core.agent.world_maintenance import WorldBibleMaintenanceService
        repo = AgentRepository(ctx.novel_manager.get_workspace(title))
        pending = [asdict(item) for item in repo.list_pending_change_sets()]
        sessions = [asdict(item) for item in repo.list_sessions()]
        artifacts = repo.list_artifacts()
        advice = []
        history = []
        try:
            from core.agent.advisor import WritingAdvisorService
            advice = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager).list_advice(title)
            history = WritingAdvisorService(ctx.novel_manager, None, ctx.conversation_manager).list_history(title)
        except Exception:
            pass
        maintenance = WorldBibleMaintenanceService(ctx.novel_manager).list_pending(title)
        return {"profiles": [asdict(item) for item in AGENT_PROFILES.values()], "sessions": sessions, "pending_changes": pending, "artifacts": artifacts[:50], "advice": advice, "advisor_history": history, "pending_world_maintenance": maintenance}

    @app.get("/api/books/{title}/agent/artifacts/{artifact_id}", tags=["agent"])
    def agent_artifact_detail(title: str, artifact_id: str, ctx: WebUserContext = Depends(current_context)):
        from core.agent.repository import AgentRepository
        artifact_name = safe_name(os.path.basename(str(artifact_id or "")))
        if not artifact_name:
            raise HTTPException(status_code=404, detail="Agent 产物不存在")
        artifact = AgentRepository(ctx.novel_manager.get_workspace(title)).load_artifact(artifact_name)
        if artifact is None:
            raise HTTPException(status_code=404, detail="Agent 产物不存在")
        return {"artifact": artifact}

    @app.post("/api/books/{title}/agent/world/maintenance/{maintenance_task_id}/retry", tags=["agent"])
    def retry_world_maintenance(title: str, maintenance_task_id: str, ctx: WebUserContext = Depends(current_context)):
        _api, client, _model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.world_maintenance import WorldBibleMaintenanceService
            handle.progress("重试世界书维护", percent=10, stage="准备")
            result = WorldBibleMaintenanceService(ctx.novel_manager).retry(client, title, maintenance_task_id)
            data = asdict(result)
            handle.progress(
                "世界书维护重试完成" if result.status == "completed" else "世界书维护仍待处理",
                percent=100 if result.status == "completed" else 95,
                stage=result.status,
                data={"result": data},
            )
            if result.status != "completed":
                raise RuntimeError(result.error or "世界书维护仍待处理")
            return data
        return {"task_id": runtime.start_task(ctx.username, f"重试世界书维护《{title}》", target, metadata={"kind": "agent_world_maintenance_retry", "book": title, "maintenance_task_id": maintenance_task_id}, retryable=True)}
    @app.post("/api/books/{title}/agent/chapter/generate", tags=["agent"])
    def agent_chapter_generate(title: str, payload: AgentChapterGenerateRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)

        def target(handle):
            from core.agent.chapter_generation import AgentChapterGenerationService, AgentChapterPlan, AgentChapterRequest
            from core.agent.supervision_agent import AgentSupervisionService, SupervisionRequest
            from core.agent.world_bible_agent import WorldBibleAgentService
            from core.app_services import ChapterGenerationService
            from strategies.novel_strategy import NovelStrategy
            from utils.prompts import Prompts

            manager = ctx.novel_manager
            workspace = manager.get_workspace(title)
            record = workspace.storage.read_json(
                f"{workspace.agent_root}/chapter_runs/{payload.plan_id}.json", default={}
            ) or {}
            if not record:
                raise RuntimeError("Agent 章节计划不存在")

            auxiliary_model = auxiliary_generation_model(ctx.settings, model)
            body_model = body_generation_model(ctx.settings, model)
            req_data = dict(record.get("request") or {})
            req_data["model"] = auxiliary_model
            request = AgentChapterRequest(**req_data)
            plan = AgentChapterPlan.from_dict(dict(record.get("plan") or {}))
            if payload.candidate_id:
                candidate = next(
                    (item for item in plan.candidate_plans if item.get("candidate_id") == payload.candidate_id),
                    None,
                )
                if candidate is None:
                    raise RuntimeError("所选 Agent 章节方案不存在")
                plan = AgentChapterPlan.from_dict(candidate)
                plan.candidate_plans = list((record.get("plan") or {}).get("candidate_plans") or [])
                plan.recommended_candidate_id = str((record.get("plan") or {}).get("recommended_candidate_id") or "")

            service = AgentChapterGenerationService(
                manager,
                client,
                skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)),
            )
            result = service.generate(request, plan)
            meta = manager.load_meta(title)
            strategy = NovelStrategy()
            strategy.novel_title = title
            strategy.chapter_title = request.chapter_title
            strategy.protagonist_bio = meta.protagonist_bio
            strategy.background_story = meta.background_story
            strategy.writing_demand = meta.writing_demand
            strategy.genre = meta.genre
            strategy.style_tone = meta.style_tone
            strategy.xp_mode = bool(meta.xp_mode)
            strategy.chapter_mode = True
            params = generation_params(ctx.settings, api_config)
            body_params = {
                **params,
                "model": body_model,
                "max_tokens": max(request.target_words * 2, params["max_tokens"]),
            }
            messages = [{"role": "system", "content": Prompts.NOVEL_CHAPTER_WRITING}]
            if strategy.xp_mode:
                messages.append({"role": "system", "content": Prompts.XP_MODE_SYSTEM})
            messages.append({
                "role": "system",
                "content": (
                    f"【本章硬性字数要求】本章字数不少于{request.target_words}字。"
                    "请优先写全必要行动过程、人物选择及其后果、带目的和阻力的对话，"
                    "以及影响后文的场景变化；不得用装饰性描写、重复反应或旁白解释凑字数。"
                ),
            })
            messages.extend(strategy.build_system_messages())
            messages.append({"role": "user", "content": result.prompt})
            handle.progress("Agent 生成正文", percent=20, stage="Agent 正文")
            content = runtime._stream_completion(handle, client, messages, body_params)
            if not content.strip():
                raise RuntimeError("模型未返回章节正文")

            def supervision_progress(stage: str) -> None:
                labels = {
                    "audit": "正在审查计划、硬性要求和连续性",
                    "repair": "正在最小化修复未通过项",
                    "reaudit": "正在复检修复稿",
                }
                handle.progress(labels.get(stage, "正在监督章节"), percent=58, stage="章节监督")

            try:
                supervised = AgentSupervisionService(
                    manager,
                    lambda _action: client,
                    skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)),
                ).supervise(
                    SupervisionRequest(
                        book_title=title,
                        chapter_num=request.chapter_num,
                        chapter_title=request.chapter_title,
                        chapter_content=content,
                        chapter_outline=request.plot,
                        requirements=request.requirement or meta.writing_demand,
                        continuity_context=result.prompt,
                        target_words=request.target_words,
                        model=auxiliary_model,
                        global_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                        xp_mode=strategy.xp_mode,
                    ),
                    progress=supervision_progress,
                )
                content = supervised.content
                supervision_report = supervised.report
            except Exception as exc:
                supervision_report = {
                    "status": "warning", "audit_failed": True, "error": str(exc),
                    "outline_items": [], "hard_constraint_issues": [],
                    "continuity_issues": [], "style_issues": [], "repair_rounds": 0,
                }
                handle.progress(
                    f"章节监督跳过，保留当前正文：{exc}", percent=60, stage="章节监督"
                )
            app_service = ChapterGenerationService(manager)
            target_info = manager.get_active_generation_target(title)
            if int(target_info.get("chapter_num") or 0) != request.chapter_num:
                raise RuntimeError("章节计划已过期：活跃路径已变化，请重新规划")
            _path, saved_version = app_service.persist_chapter(
                title=title,
                chapter_num=request.chapter_num,
                chapter_title=request.chapter_title,
                content=content,
                version=int(target_info.get("version") or manager.get_next_version(title, request.chapter_num)),
                parent_id=target_info.get("parent_id"),
                prompt=result.prompt,
                model=body_model,
                temperature=body_params["temperature"],
                top_p=body_params["top_p"],
                max_tokens=body_params["max_tokens"],
                frequency_penalty=body_params["frequency_penalty"],
                requirement=request.requirement,
                plot=request.plot,
                supervision_report=supervision_report,
                agent_data={"enabled": True, "run_id": plan.plan_id, "plan": plan.to_dict(), "context_report": result.context_report},
                generation_mode="agent",
                agent_run_id=payload.plan_id,
            )
            handle.progress("生成摘要", percent=70, stage="摘要")
            summary = manager.generate_summary(
                client,
                content,
                request.chapter_num,
                request.chapter_title,
                model=auxiliary_model,
                global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                xp_mode=strategy.xp_mode,
                raise_on_error=False,
            )
            if summary.strip():
                manager.set_chapter_node_summary(title, request.chapter_num, saved_version, summary)
            manager.rebuild_plot_summary_from_tree(title)

            try:
                director_result = service.update_director_state(
                    request, plan, summary or plan.chapter_goal, auxiliary_model
                )
                manager.update_generation_record(
                    title, request.chapter_num, saved_version, story_director=director_result
                )
            except Exception as exc:
                handle.progress(
                    f"卷级导演复盘跳过：{exc}", percent=80, stage="卷级导演"
                )
            handle.progress("更新世界书", percent=84, stage="世界书")
            maintenance = WorldBibleAgentService(manager).analyze_chapter(
                client,
                title,
                request.chapter_num,
                saved_version,
                model=auxiliary_model,
                global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                xp_mode=strategy.xp_mode,
                plan=plan.to_dict(),
            )
            manager.update_generation_record(
                title, request.chapter_num, saved_version, world_maintenance=asdict(maintenance)
            )
            try:
                snapshot = app_service.create_auto_snapshot(title, request.chapter_num, saved_version)
                snapshot_id = snapshot.snapshot_id
            except Exception as exc:
                snapshot_id = ""
                handle.progress(f"Snapshot skipped: {exc}", percent=94, stage="Snapshot")
            run_record = workspace.storage.read_json(
                f"{workspace.agent_root}/chapter_runs/{payload.plan_id}.json", default={}
            ) or {}
            run_record.update({
                "status": "completed_with_pending_maintenance" if maintenance.status != "completed" else "completed",
                "chapter_num": request.chapter_num,
                "version": saved_version,
                "chapter_node_id": manager._node_id(request.chapter_num, saved_version),
                "world_maintenance_task_id": maintenance.task_id,
                "world_maintenance_status": maintenance.status,
                "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
            })
            workspace.storage.write_json(f"{workspace.agent_root}/chapter_runs/{payload.plan_id}.json", run_record)
            result_data = {
                "chapter_num": request.chapter_num,
                "version": saved_version,
                "snapshot_id": snapshot_id,
                "world_maintenance_status": maintenance.status,
            }
            handle.progress("Agent 章节生成完成", percent=100, stage="完成", data={"result": result_data})
            return result_data

        return {"task_id": runtime.start_task(ctx.username, f"Agent 生成《{title}》章节", target, metadata={"kind": "agent_chapter_generate", "book": title, "plan_id": payload.plan_id}, retryable=True)}
    @app.post("/api/books/{title}/agent/polish/plan", tags=["agent"])
    def agent_polish_plan(title: str, payload: AgentPolishPlanRequest, ctx: WebUserContext = Depends(current_context)):
        _api, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.chapter_polish import AgentChapterPolishService, AgentPolishRequest
            meta = ctx.novel_manager.ensure_chapter_tree(title)
            node = meta.chapter_nodes.get(payload.node_id) or {}
            if not node or node.get("virtual"):
                raise RuntimeError("请选择有效章节节点")
            request = AgentPolishRequest(title, payload.node_id, int(node.get("chapter_num") or 0), str(node.get("title") or node.get("display_label") or "章节"), payload.requirement, model, str(ctx.settings.get("global_user_prompt") or ""))
            plan = AgentChapterPolishService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True))).prepare(request)
            handle.progress("Agent 润色方案完成", percent=100, stage="完成", data={"plan": plan.to_dict(), "rendered": plan.render()})
            return {"plan": plan.to_dict(), "rendered": plan.render()}
        return {"task_id": runtime.start_task(ctx.username, f"Agent 润色规划《{title}》", target, metadata={"kind": "agent_polish_plan", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/agent/polish/generate", tags=["agent"])
    def agent_polish_generate(title: str, payload: AgentPolishGenerateRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.chapter_polish import AgentChapterPolishService, AgentPolishPlan, AgentPolishRequest
            workspace = ctx.novel_manager.get_workspace(title)
            record = workspace.storage.read_json(f"{workspace.agent_root}/chapter_polish_runs/{payload.plan_id}.json", default={}) or {}
            if not record:
                raise RuntimeError("Agent 润色方案不存在")
            req_data = dict(record.get("request") or {})
            req_data["model"] = model
            request = AgentPolishRequest(**req_data)
            plan = AgentPolishPlan(**dict(record.get("plan") or {}))
            service = AgentChapterPolishService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)))
            prompt, original = service.build_prompt(request, plan)
            params = generation_params(ctx.settings, api_config)
            params["model"] = model
            handle.progress("Agent 润色正文", percent=30, stage="Agent 润色")
            candidate = runtime._stream_completion(handle, client, [{"role": "user", "content": prompt}], params)
            validation = service.validate_and_repair(request, plan, original, candidate)
            if not validation.passed:
                raise RuntimeError("润色保真校验未通过，已保存失败草稿")
            app_service = __import__("core.app_services", fromlist=["ChapterGenerationService"]).ChapterGenerationService(ctx.novel_manager)
            node = ctx.novel_manager.ensure_chapter_tree(title).chapter_nodes.get(request.node_id) or {}
            _path, version = app_service.persist_chapter(title=title, chapter_num=request.chapter_num, chapter_title=request.chapter_title, content=validation.content, version=ctx.novel_manager.get_next_version(title, request.chapter_num), parent_id=node.get("parent_id"), prompt=prompt, model=model, temperature=params["temperature"], top_p=params["top_p"], max_tokens=params["max_tokens"], frequency_penalty=params["frequency_penalty"], requirement=request.requirement, plot="Agent 润色", agent_data={"polish_plan": plan.to_dict(), "fidelity_report": validation.report}, generation_mode="agent-polish-web", agent_run_id=payload.plan_id)
            snapshot = app_service.create_auto_snapshot(title, request.chapter_num, version)
            service.mark_completed(request, plan, version, snapshot.snapshot_id)
            handle.progress("Agent 润色完成", percent=100, stage="完成", data={"result": {"chapter_num": request.chapter_num, "version": version, "snapshot_id": snapshot.snapshot_id}})
            return {"chapter_num": request.chapter_num, "version": version, "snapshot_id": snapshot.snapshot_id}
        return {"task_id": runtime.start_task(ctx.username, f"Agent 润色《{title}》", target, metadata={"kind": "agent_polish_generate", "book": title, "plan_id": payload.plan_id}, retryable=True)}

    @app.post("/api/books/{title}/agent/extra/plan", tags=["agent"])
    def agent_extra_plan(title: str, payload: AgentExtraPlanRequest, ctx: WebUserContext = Depends(current_context)):
        _api, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.extra_generation import AgentExtraGenerationService, AgentExtraRequest
            request = AgentExtraRequest(title, payload.extra_type, payload.start_node_id, payload.end_node_id, payload.reference_node_id, payload.title, payload.plot, payload.requirement, payload.target_words, model, payload.manual_entity_ids, str(ctx.settings.get("global_user_prompt") or ""))
            plan = AgentExtraGenerationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True))).prepare(request)
            handle.progress("Agent 番外方案完成", percent=100, stage="完成", data={"plan": plan.to_dict(), "rendered": plan.render()})
            return {"plan": plan.to_dict(), "rendered": plan.render()}
        return {"task_id": runtime.start_task(ctx.username, f"Agent 番外规划《{title}》", target, metadata={"kind": "agent_extra_plan", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/agent/extra/generate", tags=["agent"])
    def agent_extra_generate(title: str, payload: AgentExtraGenerateRequest, ctx: WebUserContext = Depends(current_context)):
        api_config, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.extra_generation import AgentExtraGenerationService, AgentExtraPlan, AgentExtraRequest
            workspace = ctx.novel_manager.get_workspace(title)
            record = workspace.storage.read_json(f"{workspace.agent_root}/extra_runs/{payload.plan_id}.json", default={}) or {}
            if not record:
                raise RuntimeError("Agent 番外方案不存在")
            req_data = dict(record.get("request") or {})
            req_data["model"] = model
            request = AgentExtraRequest(**req_data)
            plan = AgentExtraPlan(**dict(record.get("plan") or {}))
            service = AgentExtraGenerationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)))
            result = service.generate(request, plan)
            params = generation_params(ctx.settings, api_config)
            params["model"] = model
            handle.progress("Agent 生成番外正文", percent=30, stage="Agent 番外")
            content = runtime._stream_completion(handle, client, [{"role": "user", "content": result.prompt}], params)
            node = ctx.novel_manager.save_extra_node(title, run_id=payload.plan_id, extra_type=request.extra_type, chapter_title=request.title or "Agent 番外", content=content, start_node_id=request.start_node_id, end_node_id=request.end_node_id, reference_node_id=request.reference_node_id, summary="", generation_record={"prompt": result.prompt, "model": model, "agent_data": {"plan": plan.to_dict()}})
            snapshot = ctx.novel_manager.snapshot_service(title).create("Agent 番外生成完成", source="chapter")
            service.mark_completed(title, payload.plan_id, node.get("id", ""), snapshot.snapshot_id)
            handle.progress("Agent 番外完成", percent=100, stage="完成", data={"result": {"node": node, "snapshot_id": snapshot.snapshot_id}})
            return {"node": node, "snapshot_id": snapshot.snapshot_id}
        return {"task_id": runtime.start_task(ctx.username, f"Agent 番外生成《{title}》", target, metadata={"kind": "agent_extra_generate", "book": title, "plan_id": payload.plan_id}, retryable=True)}

    @app.post("/api/books/{title}/agent/world/analyze", tags=["agent"])
    def agent_world_analyze(title: str, payload: WorldDetailAnalyzeRequest, ctx: WebUserContext = Depends(current_context)):
        _api, client, model = text_client_and_model(ctx)
        def target(handle):
            from core.agent.world_bible_agent import WorldBibleAgentService, WorldDetailRequest
            request = WorldDetailRequest(title, payload.text, model, payload.source_run_id, str(ctx.settings.get("global_user_prompt") or ""))
            plan = WorldBibleAgentService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True))).analyze_user_details(request)
            result = asdict(plan)
            handle.progress("世界书变更待审批", percent=100, stage="待审批", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"Agent 世界书分析《{title}》", target, metadata={"kind": "agent_world_analyze", "book": title}, retryable=True)}

    @app.post("/api/books/{title}/agent/world/confirm-scopes", tags=["agent"])
    def confirm_world_scopes(title: str, payload: WorldScopeConfirmRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.world_bible_agent import WorldBibleAgentService
        WorldBibleAgentService(ctx.novel_manager).confirm_scopes(title, payload.change_set_id, payload.operations)
        return {"ok": True}

    @app.post("/api/books/{title}/agent/changes/approve", tags=["agent"])
    def approve_change(title: str, payload: ChangeApprovalRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.changes import ChangeSetService
        from core.agent.repository import AgentRepository
        result = ChangeSetService(ctx.novel_manager, title, AgentRepository(ctx.novel_manager.get_workspace(title))).approve(payload.change_set_id, payload.operation_ids)
        return {"change_set": asdict(result)}

    @app.post("/api/books/{title}/agent/changes/reject", tags=["agent"])
    def reject_change(title: str, payload: ChangeApprovalRequest, ctx: WebUserContext = Depends(current_context)):
        from core.agent.changes import ChangeSetService
        from core.agent.repository import AgentRepository
        result = ChangeSetService(ctx.novel_manager, title, AgentRepository(ctx.novel_manager.get_workspace(title))).reject(payload.change_set_id)
        return {"change_set": asdict(result)}
    @app.post("/api/continuation/segment", tags=["continuation"])
    def continuation_segment(payload: ContinuationSegmentRequest, ctx: WebUserContext = Depends(current_context)):
        sections = detect_sections(payload.text) or split_text_locally(payload.text, max_chars=6000)
        return {"sections": section_dicts(sections), "method": "local"}

    @app.post("/api/continuation/uploads", tags=["continuation"])
    async def continuation_uploads(request: Request, ctx: WebUserContext = Depends(current_context)):
        uploaded = await read_multipart_text_files(request)
        if not uploaded:
            raise HTTPException(status_code=400, detail="请上传 TXT、Markdown 或 HTML 文件")
        root = continuation_upload_root(ctx)
        cleanup_old_files(root)
        result = []
        for idx, item in enumerate(sort_upload_items(uploaded), start=1):
            filename = safe_upload_filename(item["filename"])
            stored = os.path.join(root, f"{int(time.time())}_{idx}_{filename}")
            with open(stored, "w", encoding="utf-8") as f:
                f.write(item["content"])
            detected = detect_sections(item["content"])
            sections = detected or [("全文", item["content"])]
            result.append({
                "filename": item["filename"],
                "stored_path": stored,
                "chars": len(item["content"]),
                "content": item["content"],
                "needs_ai": not bool(detected),
                "sections": section_dicts(sections),
            })
        return {"files": result}

    @app.post("/api/continuation/segment-agent", tags=["continuation"])
    def continuation_segment_agent(payload: ContinuationAgentSegmentRequest, ctx: WebUserContext = Depends(current_context)):
        source = (payload.text or "").strip()
        if not source:
            return {"sections": []}
        try:
            _api, client, model = text_client_and_model(ctx)
        except WebApiConfigError:
            sections = split_text_locally(source, max_chars=6000)
            return {"sections": section_dicts(sections), "fallback": True, "error": "未配置文字 API，已使用本地分段"}
        try:
            if payload.use_agent or ctx.settings.get("novel_generation_mode") == "agent":
                from core.agent.continuation import AgentContinuationService
                books = ctx.novel_manager.list_books()
                book_title = payload.title or (books[0] if books else "")
                service = AgentContinuationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True)))
                sections = service.segment_text(source, model, book_title=book_title, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""))
                return {"sections": section_dicts(sections), "method": "agent"}
            from utils.summarize import segment_by_ai
            sections = segment_by_ai(client, source, model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""))
            return {"sections": section_dicts(sections), "method": "ai"}
        except Exception as exc:
            sections = split_text_locally(source, max_chars=6000)
            return {"sections": section_dicts(sections), "fallback": True, "error": str(exc)}

    @app.post("/api/continuation/import", tags=["continuation"])
    def continuation_import(payload: ContinuationImportRequest, ctx: WebUserContext = Depends(current_context)):
        def target(handle):
            title = payload.title.strip() or "续写作品"
            if title not in ctx.novel_manager.list_books():
                ctx.novel_manager.create_book(title)
            sections = normalize_sections(payload.sections)
            for idx, section in enumerate(sections, start=1):
                if handle.cancelled:
                    raise RuntimeError("任务已取消")
                chapter_title = section["title"] or f"导入段落 {idx}"
                if section["content"].strip():
                    ctx.novel_manager.save_chapter_version(title, ctx.novel_manager.get_next_chapter_num(title), chapter_title, section["content"])
                handle.progress(f"已导入 {idx}/{len(sections)}", percent=min(95, 10 + idx * 80 // max(1, len(sections))), stage="续写导入")
            result = {"title": title, "count": len(sections)}
            save_continuation_run(
                ctx,
                title,
                "continuation_import",
                input_chars=sum(len(section.get("content") or "") for section in sections),
                input_summary={"sections": len(sections)},
                output_summary=result,
                result=result,
            )
            handle.progress("导入完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"导入续写《{payload.title}》", target, metadata={"kind": "continuation_import", "book": payload.title}, retryable=True)}

    @app.post("/api/continuation/analyze", tags=["continuation"])
    def continuation_analyze(payload: ContinuationAnalyzeRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            _api_config, client, model = text_client_and_model(ctx)
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        def target(handle):
            from core.app_services import ContinuationService
            from core.agent.continuation import AgentContinuationService
            from utils.summarize import generate_novel_settings_from_world_bible
            title = payload.title.strip() or "续写作品"
            sections = normalize_sections(payload.sections)
            if not sections and payload.source_text.strip():
                sections = section_dicts(detect_sections(payload.source_text) or split_text_locally(payload.source_text, max_chars=6000))
            if not sections:
                raise RuntimeError("没有可分析的分段")
            if title not in ctx.novel_manager.list_books():
                ctx.novel_manager.create_book(title)
            app_service = ContinuationService(ctx.novel_manager)
            imported = []
            handle.progress("保存确认分段", percent=8, stage="保存章节")
            for idx, section in enumerate(sections, start=1):
                if handle.cancelled:
                    raise RuntimeError("任务已取消")
                chapter_num = ctx.novel_manager.get_next_chapter_num(title)
                _path, version = ctx.novel_manager.save_chapter_version(title, chapter_num, section["title"] or f"导入段落 {idx}", section["content"])
                ctx.novel_manager.switch_active_node(title, ctx.novel_manager._node_id(chapter_num, version))
                imported.append({"chapter_num": chapter_num, "version": version, "title": section["title"]})
                handle.progress(f"保存章节 {idx}/{len(sections)}", percent=8 + idx * 22 // max(1, len(sections)), stage="保存章节")
            warnings = []
            meta = ctx.novel_manager.load_meta(title)
            for idx, item in enumerate(imported, start=1):
                if handle.cancelled:
                    raise RuntimeError("任务已取消")
                content = sections[idx - 1]["content"]
                try:
                    ctx.novel_manager.generate_summary(client, content, int(item["chapter_num"]), item["title"], model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(payload.xp_mode or meta.xp_mode), raise_on_error=True)
                except Exception as exc:
                    warnings.append(f"章节 {item['chapter_num']} 摘要失败：{exc}")
                handle.progress(f"生成摘要 {idx}/{len(imported)}", percent=32 + idx * 20 // max(1, len(imported)), stage="生成摘要")
            handle.progress("抽取并合并世界书", percent=58, stage="世界书")
            try:
                ctx.novel_manager.rebuild_world_bible_from_active(client, title, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(payload.xp_mode or meta.xp_mode), force_extract=True, extract_missing=True)
            except Exception as exc:
                warnings.append(f"世界书重建失败：{exc}")
            world_data = world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
            handle.progress("从世界书生成小说设定", percent=76, stage="生成设定")
            if ctx.settings.get("novel_generation_mode") == "agent":
                settings = AgentContinuationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True))).generate_settings_from_world_data(world_data, model, book_title=title, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(payload.xp_mode or meta.xp_mode))
            else:
                settings = generate_novel_settings_from_world_bible(client, world_data, model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=bool(payload.xp_mode or meta.xp_mode))
            saved_meta = ctx.novel_manager.save_meta(title, protagonist_bio=settings.get("protagonist_bio", meta.protagonist_bio), background_story=settings.get("background_story", meta.background_story), writing_demand=settings.get("writing_demand", meta.writing_demand), author_plan=settings.get("author_plan", meta.author_plan), genre=getattr(meta, "genre", ""), style_tone=getattr(meta, "style_tone", ""), xp_mode=bool(payload.xp_mode or meta.xp_mode))
            handle.progress("创建项目快照", percent=92, stage="快照")
            snapshot_id = ""
            try:
                snapshot_id = app_service.create_auto_snapshot(title, int(imported[-1]["chapter_num"]), int(imported[-1]["version"])).snapshot_id
            except Exception as exc:
                warnings.append(f"快照创建失败：{exc}")
            result = {"title": title, "imported": imported, "settings": settings, "meta": serialize_meta(saved_meta), "world_counts": world_counts(world_data), "snapshot_id": snapshot_id, "warnings": warnings}
            save_continuation_run(
                ctx,
                title,
                "continuation_analyze",
                input_chars=sum(len(section.get("content") or "") for section in sections),
                input_summary={"sections": len(sections), "source_text_chars": len(payload.source_text or ""), "xp_mode": bool(payload.xp_mode or meta.xp_mode)},
                output_summary={"imported": len(imported), "world_counts": result["world_counts"], "settings_fields": sorted(settings.keys()), "snapshot_id": snapshot_id, "warnings": len(warnings)},
                result=result,
            )
            handle.progress("分析导入完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"分析旧文并建书《{payload.title}》", target, metadata={"kind": "continuation_analyze", "book": payload.title}, retryable=True)}

    @app.post("/api/continuation/suggest", tags=["continuation"])
    def continuation_suggest(payload: ContinuationSuggestRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            _api, client, model = text_client_and_model(ctx)
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        def target(handle):
            from core.agent.continuation import AgentContinuationService
            from ui.continuation_dialogs import suggest_directions
            books = ctx.novel_manager.list_books()
            title = (payload.title or (books[0] if books else "")).strip()
            setting = payload.setting.strip()
            plot = payload.plot.strip()
            world_data = payload.world_data or {}
            if title and not setting:
                meta = ctx.novel_manager.load_meta(title)
                setting = "\n\n".join([meta.background_story, meta.protagonist_bio, meta.writing_demand, meta.author_plan]).strip()
                try:
                    world_data = world_data or world_bible_to_dict(ctx.novel_manager.load_world_bible(title))
                except Exception:
                    pass
            handle.progress("生成发展方向", percent=35, stage="方向建议")
            if ctx.settings.get("novel_generation_mode") == "agent":
                directions = AgentContinuationService(ctx.novel_manager, client, skills_enabled=bool(ctx.settings.get("agent_skills_enabled", True))).suggest_directions(setting, plot, model, book_title=title, world_data=world_data, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=payload.xp_mode)
            else:
                directions = suggest_directions(client, setting, plot, model, world_data=world_data, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=payload.xp_mode)
            result = {"title": title, "directions": directions, "setting": setting, "plot": plot}
            if title:
                save_continuation_run(
                    ctx,
                    title,
                    "continuation_direction",
                    input_chars=len(setting) + len(plot),
                    input_summary={"setting_chars": len(setting), "plot_chars": len(plot), "xp_mode": bool(payload.xp_mode)},
                    output_summary={"directions": len(directions)},
                    result=result,
                )
            handle.progress("方向建议完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, "续写发展方向建议", target, metadata={"kind": "continuation_suggest", "book": payload.title}, retryable=True)}

    @app.post("/api/continuation/generate", tags=["continuation"])
    def continuation_generate(payload: ContinuationGenerateRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            api_config, client, model = text_client_and_model(ctx)
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        def target(handle):
            from core.app_services import ContinuationService
            from strategies.continuation_strategy import ContinuationStrategy
            title = payload.title.strip() or "续写作品"
            manager = ctx.novel_manager
            chapter_mode = bool(payload.chapter_mode)
            book_exists = title in manager.list_books()
            if chapter_mode and not book_exists:
                manager.create_book(title)
                book_exists = True
            service = ContinuationService(manager)
            params = generation_params(ctx.settings, api_config)
            params["model"] = model
            meta = manager.load_meta(title) if book_exists else NovelMeta(title=title)
            if chapter_mode:
                target_info = manager.get_active_generation_target(title)
                chapter_num = int(target_info.get("chapter_num") or manager.get_next_chapter_num(title))
                parent_id = target_info.get("parent_id")
                chapter_title = payload.chapter_title.strip() or f"第{chapter_num}章"
            else:
                target_info = {}
                chapter_num = int(manager.get_next_chapter_num(title) if book_exists else 1)
                parent_id = None
                chapter_title = payload.chapter_title.strip() or "续写草稿"
            xp_mode = bool(payload.xp_mode or meta.xp_mode)
            handle.progress("准备续写上下文", percent=8, stage="准备上下文")
            if book_exists:
                context_report = service.build_context(title, chapter_num, chapter_title, payload.source_text, payload.requirement, payload.plot, global_prompt=str(ctx.settings.get("global_user_prompt") or ""), client=client, model=model)
                context_text = context_report.render()
            else:
                context_text = "当前未写入书架，仅基于源文档和用户要求生成续写草稿。"
            prompt = build_continuation_prompt(title=title, chapter_num=chapter_num, chapter_title=chapter_title, source_text=payload.source_text, requirement=payload.requirement, plot=payload.plot, setting=payload.setting, target_words=payload.target_words, meta=meta, context_text=context_text, xp_mode=xp_mode, chapter_mode=chapter_mode)
            messages = [{"role": "system", "content": ContinuationStrategy().get_system_prompt()}, {"role": "user", "content": prompt}]
            handle.progress("生成续写正文", percent=20, stage="生成正文")
            content = runtime._stream_completion(handle, client, messages, params)
            if not content.strip():
                raise RuntimeError("模型未返回续写正文")
            handle.progress("监督修补", percent=58, stage="监督修补")
            supervision_report = {"status": "not_available"}
            try:
                from utils.supervision import supervise_chapter
                def supervision_client(_kind):
                    return client
                content, report = supervise_chapter(
                    supervision_client,
                    chapter_content=content,
                    chapter_title=chapter_title,
                    chapter_outline=payload.plot,
                    requirements=payload.requirement,
                    continuity_context=prompt,
                    target_words=payload.target_words,
                    model=model,
                    temperature=params["temperature"],
                    global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""),
                    xp_mode=xp_mode,
                    max_repair_rounds=2,
                )
                supervision_report = report.to_dict()
            except Exception as exc:
                supervision_report = {"status": "warning", "error": str(exc)}
            warnings = []
            if not chapter_mode:
                result = {"title": title, "chapter_num": chapter_num, "chapter_title": chapter_title, "chapter_mode": False, "draft_only": True, "supervision_report": supervision_report, "warnings": warnings, "content": content, "preview": content[:240]}
                save_continuation_run(
                    ctx,
                    title if book_exists else "",
                    "continuation_draft",
                    input_chars=len(payload.source_text or "") + len(payload.requirement or "") + len(payload.plot or "") + len(payload.setting or ""),
                    input_summary={"source_text_chars": len(payload.source_text or ""), "requirement": payload.requirement, "plot": payload.plot, "setting": payload.setting, "target_words": payload.target_words, "chapter_mode": False, "xp_mode": xp_mode},
                    output_summary={"draft_only": True, "content_chars": len(content), "warnings": len(warnings)},
                    result=result,
                )
                handle.progress("续写草稿完成", percent=100, stage="完成", data={"result": result})
                return result
            handle.progress("保存章节版本", percent=64, stage="保存章节")
            _path, saved_version = service.persist_chapter(title=title, chapter_num=chapter_num, chapter_title=chapter_title, content=content, version=manager.get_next_version(title, chapter_num), parent_id=parent_id, prompt=prompt, model=model, temperature=params["temperature"], top_p=params["top_p"], max_tokens=params["max_tokens"], frequency_penalty=params["frequency_penalty"], supervision_report=supervision_report, requirement=payload.requirement, plot=payload.plot, generation_mode="continuation-web")
            handle.progress("生成章节摘要", percent=74, stage="生成摘要")
            try:
                manager.generate_summary(client, content, chapter_num, chapter_title, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=xp_mode, raise_on_error=True)
            except Exception as exc:
                warnings.append(f"摘要生成失败：{exc}")
            handle.progress("更新世界书", percent=84, stage="更新世界书")
            try:
                service.world_bible.sync_chapter(client, title, chapter_num, saved_version, content, model=model, global_user_prompt=str(ctx.settings.get("global_user_prompt") or ""), xp_mode=xp_mode)
            except Exception as exc:
                warnings.append(f"世界书更新失败：{exc}")
            handle.progress("创建项目快照", percent=93, stage="创建快照")
            snapshot_id = ""
            try:
                snapshot_id = service.create_auto_snapshot(title, chapter_num, saved_version).snapshot_id
            except Exception as exc:
                warnings.append(f"快照创建失败：{exc}")
            result = {"title": title, "chapter_num": chapter_num, "chapter_title": chapter_title, "chapter_mode": True, "draft_only": False, "version": saved_version, "snapshot_id": snapshot_id, "warnings": warnings, "preview": content[:240]}
            save_continuation_run(
                ctx,
                title,
                "continuation_generate",
                input_chars=len(payload.source_text or "") + len(payload.requirement or "") + len(payload.plot or "") + len(payload.setting or ""),
                input_summary={"source_text_chars": len(payload.source_text or ""), "requirement": payload.requirement, "plot": payload.plot, "setting": payload.setting, "target_words": payload.target_words, "chapter_mode": True, "xp_mode": xp_mode},
                output_summary={"chapter_num": chapter_num, "version": saved_version, "snapshot_id": snapshot_id, "warnings": len(warnings)},
                result=result,
            )
            handle.progress("续写完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"续写生成《{payload.title}》", target, metadata={"kind": "continuation_generate", "book": payload.title}, retryable=True)}

    @app.get("/api/continuation/runs", tags=["continuation"])
    def continuation_runs(title: str = Query(default=""), ctx: WebUserContext = Depends(current_context)):
        return {"runs": list_continuation_runs(ctx, title)}

    @app.get("/api/continuation/runs/{run_id}", tags=["continuation"])
    def continuation_run_detail(run_id: str, title: str = Query(default=""), ctx: WebUserContext = Depends(current_context)):
        run = load_continuation_run(ctx, run_id, title=title)
        if not run:
            raise HTTPException(status_code=404, detail="续写运行记录不存在")
        return {"run": run}

    @app.get("/api/markdown/tree", tags=["markdown"])
    def markdown_tree(ctx: WebUserContext = Depends(current_context)):
        return {"items": markdown_tree_items(ctx)}

    @app.post("/api/markdown/folder", tags=["markdown"])
    def create_markdown_folder(payload: MarkdownFolderRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            path = ctx.markdown_path(payload.path)
            os.makedirs(path, exist_ok=False)
            return {"ok": True, "path": markdown_display_path(ctx, path)}
        except FileExistsError as exc:
            raise HTTPException(status_code=409, detail="文件夹已存在") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.get("/api/markdown/file", tags=["markdown"])
    def read_markdown(path: str, ctx: WebUserContext = Depends(current_context)):
        try:
            return {"path": path, "content": ctx.read_markdown(path)}
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail="文件不存在") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.get("/api/markdown/preview", tags=["markdown"])
    def preview_markdown(path: str, ctx: WebUserContext = Depends(current_context)):
        try:
            content = ctx.read_markdown(path)
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail="文件不存在") from exc
        try:
            import markdown as md_lib
            body = md_lib.markdown(content, extensions=["fenced_code", "tables", "nl2br", "sane_lists"])
        except Exception:
            body = f"<pre>{html.escape(content)}</pre>"
        return {"path": path, "html": f"<article class='markdown-preview'>{body}</article>"}

    @app.put("/api/markdown/file", tags=["markdown"])
    def write_markdown(payload: MarkdownWriteRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            stored = ctx.write_markdown(payload.path, payload.content)
            return {"ok": True, "path": markdown_display_path(ctx, stored)}
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.post("/api/markdown/rename", tags=["markdown"])
    def rename_markdown(payload: MarkdownRenameRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            source = ctx.markdown_storage_path(payload.path)
            if not os.path.exists(source):
                source = ctx.markdown_path(payload.path)
            if not os.path.exists(source):
                raise FileNotFoundError(payload.path)
            if os.path.abspath(source) == os.path.abspath(ctx.markdown_root):
                raise ValueError("不能重命名笔记根目录")
            destination = ctx.markdown_path(payload.new_path) if os.path.isdir(source) else ctx.markdown_storage_path(payload.new_path, for_write=True)
            if os.path.exists(destination):
                raise FileExistsError(payload.new_path)
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            os.replace(source, destination)
            return {"ok": True, "path": markdown_display_path(ctx, destination)}
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail="文件或文件夹不存在") from exc
        except FileExistsError as exc:
            raise HTTPException(status_code=409, detail="目标已存在") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.delete("/api/markdown/path", tags=["markdown"])
    def delete_markdown_path(path: str, ctx: WebUserContext = Depends(current_context)):
        try:
            target = ctx.markdown_storage_path(path)
            if not os.path.exists(target):
                target = ctx.markdown_path(path)
            if not os.path.exists(target):
                raise FileNotFoundError(path)
            if os.path.abspath(target) == os.path.abspath(ctx.markdown_root):
                raise ValueError("不能删除笔记根目录")
            if os.path.isdir(target):
                shutil.rmtree(target)
            else:
                os.remove(target)
            return {"ok": True}
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail="文件或文件夹不存在") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.post("/api/markdown/export", tags=["markdown"])
    def export_markdown(payload: MarkdownExportRequest, ctx: WebUserContext = Depends(current_context)):
        runtime.cleanup_exports(ctx)
        try:
            if payload.folder:
                source = ctx.markdown_path(payload.path) if payload.path else os.path.abspath(ctx.markdown_root)
                if not os.path.isdir(source):
                    raise FileNotFoundError(payload.path)
                name = safe_name(os.path.basename(source.rstrip(os.sep)) or "markdown_notes")
                out = os.path.join(ctx.export_root, f"{name}.zip")
                with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as archive:
                    root = os.path.abspath(source)
                    for current, _dirs, files in os.walk(root):
                        for filename in sorted(files):
                            full = os.path.join(current, filename)
                            if not markdown_is_note_file(ctx, full):
                                continue
                            display = markdown_display_path(ctx, full)
                            rel = os.path.relpath(os.path.join(ctx.markdown_root, display), root).replace("\\", "/")
                            archive.writestr(rel, ctx.read_markdown(display))
                download = runtime.register_download(ctx.username, out, os.path.basename(out), media_type_for(out))
                return {"download": download}
            content = ctx.read_markdown(payload.path)
            stem = safe_name(os.path.splitext(os.path.basename(payload.path.replace(".enc", "")))[0] or "note")
            out = os.path.join(ctx.export_root, f"{stem}.md")
            Path(out).write_text(content, encoding="utf-8")
            download = runtime.register_download(ctx.username, out, os.path.basename(out), media_type_for(out))
            return {"download": download}
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail="文件或文件夹不存在") from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    @app.get("/api/roleplay/characters", tags=["roleplay"])
    def list_characters(ctx: WebUserContext = Depends(current_context)):
        return {"book": character_book_to_dict(ctx.character_book_manager.load())}

    @app.post("/api/roleplay/characters", tags=["roleplay"])
    def create_character(payload: RoleProfileRequest, ctx: WebUserContext = Depends(current_context)):
        profile = CharacterProfile(**filter_profile(payload.profile))
        created = ctx.character_book_manager.create_profile(profile)
        return {"profile": asdict(created)}

    @app.put("/api/roleplay/characters/{character_id}", tags=["roleplay"])
    def update_character(character_id: str, payload: RoleProfileRequest, ctx: WebUserContext = Depends(current_context)):
        data = dict(payload.profile or {})
        data["character_id"] = character_id
        profile = CharacterProfile(**filter_profile(data))
        ctx.character_book_manager.update_profile(profile)
        return {"profile": asdict(profile)}

    @app.delete("/api/roleplay/characters/{character_id}", tags=["roleplay"])
    def delete_character(character_id: str, ctx: WebUserContext = Depends(current_context)):
        ctx.character_book_manager.delete_profile(character_id)
        return {"ok": True}

    @app.get("/api/roleplay/character-book", tags=["roleplay"])
    def read_character_book(ctx: WebUserContext = Depends(current_context)):
        return {"book": character_book_to_dict(ctx.character_book_manager.load())}

    @app.put("/api/roleplay/character-book", tags=["roleplay"])
    def save_character_book(payload: CharacterBookRequest, ctx: WebUserContext = Depends(current_context)):
        book = dict_to_character_book(payload.book or {})
        ctx.character_book_manager.save(book)
        return {"book": character_book_to_dict(ctx.character_book_manager.load())}

    @app.get("/api/roleplay/conversations/{conversation_id}/memory", tags=["roleplay"])
    def read_conversation_memory(conversation_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        return roleplay_memory_state(ctx, record)

    @app.post("/api/roleplay/conversations/{conversation_id}/memory/{change_set_id}/apply", tags=["roleplay"])
    def apply_conversation_memory_change(conversation_id: str, change_set_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        idx, change_set = find_memory_change_set(record, change_set_id)
        if change_set.status == "applied":
            return roleplay_memory_state(ctx, record)
        book = ctx.character_book_manager.load()
        apply_memory_change_set(book, change_set)
        ctx.character_book_manager.save(book)
        update_memory_change_set(record, idx, change_set)
        snapshot_roleplay_character_book(record, book)
        save_roleplay_record(ctx, record)
        return roleplay_memory_state(ctx, ctx.conversation_manager.load_conversation(conversation_id) or record)

    @app.post("/api/roleplay/conversations/{conversation_id}/memory/{change_set_id}/reject", tags=["roleplay"])
    def reject_conversation_memory_change(conversation_id: str, change_set_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        idx, change_set = find_memory_change_set(record, change_set_id)
        change_set.status = "rejected"
        update_memory_change_set(record, idx, change_set)
        save_roleplay_record(ctx, record)
        return roleplay_memory_state(ctx, ctx.conversation_manager.load_conversation(conversation_id) or record)

    @app.post("/api/roleplay/conversations/{conversation_id}/memory/{change_set_id}/revert", tags=["roleplay"])
    def revert_conversation_memory_change(conversation_id: str, change_set_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        idx, change_set = find_memory_change_set(record, change_set_id)
        if change_set.status != "applied":
            raise HTTPException(status_code=400, detail="只有已应用的记忆变更可以撤销")
        book = ctx.character_book_manager.load()
        revert_memory_change_set(book, change_set)
        ctx.character_book_manager.save(book)
        update_memory_change_set(record, idx, change_set)
        snapshot_roleplay_character_book(record, book)
        save_roleplay_record(ctx, record)
        return roleplay_memory_state(ctx, ctx.conversation_manager.load_conversation(conversation_id) or record)

    @app.put("/api/roleplay/conversations/{conversation_id}/memory/{change_set_id}", tags=["roleplay"])
    def edit_conversation_memory_change(conversation_id: str, change_set_id: str, payload: MemoryChangeEditRequest, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        idx, change_set = find_memory_change_set(record, change_set_id)
        update_by_id = {str(item.get("change_id")): item for item in payload.changes if isinstance(item, dict) and item.get("change_id")}
        for change in change_set.changes:
            update = update_by_id.get(change.change_id)
            if not update:
                continue
            if "new_value" in update:
                change.new_value = update.get("new_value")
            if "reason" in update:
                change.reason = str(update.get("reason", ""))
            if "risk" in update:
                change.risk = str(update.get("risk", change.risk))
        update_memory_change_set(record, idx, change_set)
        save_roleplay_record(ctx, record)
        if payload.apply_now:
            return apply_conversation_memory_change(conversation_id, change_set_id, ctx)
        return roleplay_memory_state(ctx, ctx.conversation_manager.load_conversation(conversation_id) or record)
    @app.get("/api/roleplay/senders", tags=["roleplay"])
    def list_sender_profiles(ctx: WebUserContext = Depends(current_context)):
        return {"profiles": [asdict(item) for item in ctx.sender_profile_manager.load()]}

    @app.post("/api/roleplay/senders", tags=["roleplay"])
    def create_sender_profile(payload: SenderProfileRequest, ctx: WebUserContext = Depends(current_context)):
        profile = SenderProfile(**filter_fields(SenderProfile, payload.profile))
        profiles = ctx.sender_profile_manager.load()
        profiles.append(profile)
        ctx.sender_profile_manager.save(profiles)
        return {"profile": asdict(profile)}

    @app.put("/api/roleplay/senders/{profile_id}", tags=["roleplay"])
    def update_sender_profile(profile_id: str, payload: SenderProfileRequest, ctx: WebUserContext = Depends(current_context)):
        profiles = ctx.sender_profile_manager.load()
        for index, profile in enumerate(profiles):
            if profile.sender_profile_id == profile_id:
                data = dict(payload.profile or {})
                data["sender_profile_id"] = profile_id
                data.setdefault("created_at", profile.created_at)
                profiles[index] = SenderProfile(**filter_fields(SenderProfile, data))
                ctx.sender_profile_manager.save(profiles)
                return {"profile": asdict(profiles[index])}
        raise HTTPException(status_code=404, detail="发送者档案不存在")

    @app.delete("/api/roleplay/senders/{profile_id}", tags=["roleplay"])
    def delete_sender_profile(profile_id: str, ctx: WebUserContext = Depends(current_context)):
        profiles = ctx.sender_profile_manager.load()
        kept = [profile for profile in profiles if profile.sender_profile_id != profile_id]
        if len(kept) == len(profiles):
            raise HTTPException(status_code=404, detail="发送者档案不存在")
        ctx.sender_profile_manager.save(kept)
        return {"ok": True}

    @app.get("/api/roleplay/scenes", tags=["roleplay"])
    def list_scene_presets(ctx: WebUserContext = Depends(current_context)):
        return {"presets": [asdict(item) for item in ctx.scene_preset_manager.load()]}

    @app.post("/api/roleplay/scenes", tags=["roleplay"])
    def create_scene_preset(payload: ScenePresetRequest, ctx: WebUserContext = Depends(current_context)):
        data = dict(payload.preset or {})
        scene = SceneState(**filter_fields(SceneState, data.get("scene") or {}))
        preset = ScenePreset(scene_preset_id=data.get("scene_preset_id", ""), name=data.get("name", ""), scene=scene, created_at=data.get("created_at", ""), updated_at=data.get("updated_at", ""))
        presets = ctx.scene_preset_manager.load()
        presets.append(preset)
        ctx.scene_preset_manager.save(presets)
        return {"preset": asdict(preset)}

    @app.put("/api/roleplay/scenes/{preset_id}", tags=["roleplay"])
    def update_scene_preset(preset_id: str, payload: ScenePresetRequest, ctx: WebUserContext = Depends(current_context)):
        presets = ctx.scene_preset_manager.load()
        for index, preset in enumerate(presets):
            if preset.scene_preset_id == preset_id:
                data = dict(payload.preset or {})
                scene = SceneState(**filter_fields(SceneState, data.get("scene") or {}))
                presets[index] = ScenePreset(scene_preset_id=preset_id, name=data.get("name", preset.name), scene=scene, created_at=preset.created_at, updated_at=data.get("updated_at", ""))
                ctx.scene_preset_manager.save(presets)
                return {"preset": asdict(presets[index])}
        raise HTTPException(status_code=404, detail="场景预设不存在")

    @app.delete("/api/roleplay/scenes/{preset_id}", tags=["roleplay"])
    def delete_scene_preset(preset_id: str, ctx: WebUserContext = Depends(current_context)):
        presets = ctx.scene_preset_manager.load()
        kept = [preset for preset in presets if preset.scene_preset_id != preset_id]
        if len(kept) == len(presets):
            raise HTTPException(status_code=404, detail="场景预设不存在")
        ctx.scene_preset_manager.save(kept)
        return {"ok": True}

    @app.get("/api/roleplay/conversations/{conversation_id}/controls", tags=["roleplay"])
    def read_conversation_controls(conversation_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        return {"state": roleplay_control_state(record), "sender_profiles": [asdict(item) for item in ctx.sender_profile_manager.load()], "scene_presets": [asdict(item) for item in ctx.scene_preset_manager.load()]}

    @app.put("/api/roleplay/conversations/{conversation_id}/controls", tags=["roleplay"])
    def save_conversation_controls(conversation_id: str, payload: ChatControlRequest, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        state = payload.state or {}
        scene = SceneState(**filter_fields(SceneState, state.get("scene_state") or record.get("scene_state") or {}))
        turn_policy = TurnPolicy(**filter_fields(TurnPolicy, state.get("turn_policy") or record.get("turn_policy") or {}))
        ctx.conversation_manager.save_conversation(
            conversation_id=conversation_id,
            title=record.get("title") or "角色对话",
            model=record.get("model") or "",
            messages=record.get("messages") or [],
            character_description=record.get("character_description") or "",
            story_background=record.get("story_background") or "",
            strategy=record.get("strategy") or "角色扮演",
            reply_mode=state.get("reply_mode") or record.get("reply_mode") or "character",
            chat_type=state.get("chat_type") or record.get("chat_type") or "private",
            participant_character_ids=state.get("participant_character_ids") or record.get("participant_character_ids") or [],
            primary_character_id=record.get("primary_character_id") or "",
            timeline_id=record.get("timeline_id") or conversation_id,
            timeline=record.get("timeline") or [],
            character_book_snapshot=record.get("character_book_snapshot") or {},
            sender_name=state.get("sender_name") or record.get("sender_name") or "你",
            sender_profile=state.get("sender_profile") or record.get("sender_profile") or "",
            required_responder_ids=state.get("required_responder_ids") or record.get("required_responder_ids") or [],
            structured_messages=record.get("structured_messages") or [],
            branches=record.get("branches") or [],
            active_branch_id=state.get("active_branch_id") or record.get("active_branch_id") or "main",
            sender_profile_id=state.get("sender_profile_id") or record.get("sender_profile_id") or "",
            scene_state=asdict(scene),
            turn_policy=asdict(turn_policy),
            memory_change_sets=record.get("memory_change_sets") or [],
            narrator_enabled=bool(state.get("narrator_enabled", record.get("narrator_enabled", False))),
            schema_version=int(record.get("schema_version") or 1),
        )
        updated = ctx.conversation_manager.load_conversation(conversation_id) or {}
        return {"state": roleplay_control_state(updated)}

    @app.get("/api/roleplay/conversations", tags=["roleplay"])
    def list_conversations(ctx: WebUserContext = Depends(current_context)):
        return {"conversations": [asdict(item) for item in ctx.conversation_manager.list_conversations()]}

    @app.get("/api/roleplay/conversations/{conversation_id}", tags=["roleplay"])
    def read_conversation(conversation_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        return {"conversation": record}

    @app.delete("/api/roleplay/conversations/{conversation_id}", tags=["roleplay"])
    def delete_conversation(conversation_id: str, ctx: WebUserContext = Depends(current_context)):
        if not ctx.conversation_manager.delete_conversation(conversation_id):
            raise HTTPException(status_code=404, detail="会话不存在")
        return {"ok": True}

    @app.get("/api/roleplay/conversations/{conversation_id}/branches", tags=["roleplay"])
    def list_conversation_branches(conversation_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        normalized = ensure_roleplay_branches(record)
        if normalized != record:
            save_roleplay_record(ctx, normalized)
        return {"branches": normalized.get("branches") or [], "active_branch_id": normalized.get("active_branch_id") or "main"}

    @app.post("/api/roleplay/conversations/{conversation_id}/branches/fork", tags=["roleplay"])
    def fork_conversation_branch(conversation_id: str, payload: ConversationBranchRequest, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        active_id = record.get("active_branch_id") or "main"
        parent = next((item for item in record.get("branches", []) if item.get("branch_id") == active_id), None)
        if not parent:
            raise HTTPException(status_code=404, detail="活跃分支不存在")
        messages = list(parent.get("messages") or [])
        fork_id = payload.message_id.strip()
        if fork_id:
            cut = next((idx for idx, item in enumerate(messages) if item.get("message_id") == fork_id), -1)
            if cut < 0:
                raise HTTPException(status_code=404, detail="分叉消息不存在")
            messages = messages[:cut + 1]
        else:
            fork_id = messages[-1].get("message_id", "") if messages else ""
        from core.chat_domain import new_id, now_text
        branch = {
            "branch_id": new_id("branch"),
            "title": payload.title.strip() or f"分支 {len(record.get('branches', [])) + 1}",
            "parent_branch_id": active_id,
            "fork_message_id": fork_id,
            "messages": messages,
            "timeline": list(parent.get("timeline") or record.get("timeline") or []),
            "character_state_snapshot": dict(parent.get("character_state_snapshot") or {}),
            "knowledge": list(parent.get("knowledge") or []),
            "relationships": list(parent.get("relationships") or []),
            "created_at": now_text(),
        }
        record.setdefault("branches", []).append(branch)
        record["active_branch_id"] = branch["branch_id"]
        record["structured_messages"] = list(branch["messages"])
        save_roleplay_record(ctx, record)
        return {"branch": branch, "branches": record.get("branches") or [], "active_branch_id": record["active_branch_id"]}

    @app.post("/api/roleplay/conversations/{conversation_id}/branches/{branch_id}/activate", tags=["roleplay"])
    def activate_conversation_branch(conversation_id: str, branch_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        branch = next((item for item in record.get("branches", []) if item.get("branch_id") == branch_id), None)
        if not branch:
            raise HTTPException(status_code=404, detail="分支不存在")
        record["active_branch_id"] = branch_id
        record["structured_messages"] = list(branch.get("messages") or [])
        record["timeline"] = list(branch.get("timeline") or record.get("timeline") or [])
        save_roleplay_record(ctx, record)
        return {"conversation": record}

    @app.delete("/api/roleplay/conversations/{conversation_id}/branches/{branch_id}", tags=["roleplay"])
    def delete_conversation_branch(conversation_id: str, branch_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        if branch_id == "main":
            raise HTTPException(status_code=400, detail="主线分支不能删除")
        branches = list(record.get("branches") or [])
        kept = [item for item in branches if item.get("branch_id") != branch_id]
        if len(kept) == len(branches):
            raise HTTPException(status_code=404, detail="分支不存在")
        record["branches"] = kept
        if record.get("active_branch_id") == branch_id:
            active = kept[0]
            record["active_branch_id"] = active.get("branch_id") or "main"
            record["structured_messages"] = list(active.get("messages") or [])
        save_roleplay_record(ctx, record)
        return {"ok": True, "branches": record.get("branches") or [], "active_branch_id": record.get("active_branch_id") or "main"}


    @app.get("/api/roleplay/conversations/{conversation_id}/messages/{message_id}", tags=["roleplay"])
    def read_conversation_message(conversation_id: str, message_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        branch = active_roleplay_branch(record)
        message = next((item for item in branch.get("messages", []) if item.get("message_id") == message_id), None)
        if not message:
            raise HTTPException(status_code=404, detail="消息不存在")
        related = [item for item in record.get("memory_change_sets") or [] if message_id in (item.get("source_message_ids") or [])]
        return {"message": message, "source": {"message_id": message.get("message_id"), "source_message_id": message.get("source_message_id", ""), "branch_id": message.get("branch_id", branch.get("branch_id", "main")), "turn_index": message.get("turn_index", 0)}, "memory_change_sets": related}

    @app.put("/api/roleplay/conversations/{conversation_id}/messages/{message_id}", tags=["roleplay"])
    def edit_conversation_message(conversation_id: str, message_id: str, payload: ConversationMessageRequest, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        branch = active_roleplay_branch(record)
        messages = list(branch.get("messages") or [])
        index = next((idx for idx, item in enumerate(messages) if item.get("message_id") == message_id), -1)
        if index < 0:
            raise HTTPException(status_code=404, detail="消息不存在")
        messages[index]["content"] = payload.content.strip()
        branch["messages"] = messages
        record["structured_messages"] = list(messages)
        save_roleplay_record(ctx, record)
        return {"conversation": ctx.conversation_manager.load_conversation(conversation_id) or record}

    @app.delete("/api/roleplay/conversations/{conversation_id}/messages/{message_id}", tags=["roleplay"])
    def delete_conversation_message(conversation_id: str, message_id: str, ctx: WebUserContext = Depends(current_context)):
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        record = ensure_roleplay_branches(record)
        branch = active_roleplay_branch(record)
        messages = list(branch.get("messages") or [])
        kept = [item for item in messages if item.get("message_id") != message_id]
        if len(kept) == len(messages):
            raise HTTPException(status_code=404, detail="消息不存在")
        branch["messages"] = kept
        record["structured_messages"] = list(kept)
        save_roleplay_record(ctx, record)
        return {"conversation": ctx.conversation_manager.load_conversation(conversation_id) or record}

    @app.post("/api/roleplay/conversations/{conversation_id}/messages/{message_id}/fork", tags=["roleplay"])
    def fork_conversation_message(conversation_id: str, message_id: str, payload: ConversationMessageRequest, ctx: WebUserContext = Depends(current_context)):
        return fork_conversation_branch(conversation_id, ConversationBranchRequest(message_id=message_id, title=payload.title), ctx)

    @app.post("/api/roleplay/conversations/{conversation_id}/messages/{message_id}/regenerate", tags=["roleplay"])
    def regenerate_conversation_message(conversation_id: str, message_id: str, payload: ConversationMessageRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            api_config, client, model = text_client_and_model(ctx)
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        def target(handle):
            from core.chat_domain import ChatMessage, parse_structured_reply, structured_to_legacy_messages, now_text, new_id
            from strategies.role_play_strategy import RolePlayStrategy
            record = ctx.conversation_manager.load_conversation(conversation_id)
            if not record:
                raise RuntimeError("会话不存在")
            record = ensure_roleplay_branches(record)
            branch = active_roleplay_branch(record)
            messages = list(branch.get("messages") or [])
            index = next((idx for idx, item in enumerate(messages) if item.get("message_id") == message_id), -1)
            if index < 0:
                raise RuntimeError("消息不存在")
            original = messages[index]
            if original.get("role") != "assistant":
                raise RuntimeError("只能重生成角色发言")
            fork_at = messages[index - 1].get("message_id", "") if index > 0 else message_id
            fork_payload = ConversationBranchRequest(message_id=fork_at, title=payload.title or f"重生成-{original.get('speaker_name') or '角色'}")
            branch_data = fork_conversation_branch(conversation_id, fork_payload, ctx)["branch"]
            branch_messages = [ChatMessage(**filter_chat_message(item)) for item in branch_data.get("messages") or []]
            speaker_name = original.get("speaker_name") or "角色"
            speaker_id = original.get("speaker_id") or "assistant"
            strategy = RolePlayStrategy()
            book = ctx.character_book_manager.load()
            strategy.character_book = book
            strategy.participant_character_ids = [speaker_id]
            strategy.primary_character_id = speaker_id
            strategy.sender_name = record.get("sender_name") or "你"
            strategy.sender_profile = record.get("sender_profile") or ""
            legacy = structured_to_legacy_messages(branch_messages, strategy.get_system_prompt())
            legacy.append({"role": "user", "content": f"只让角色「{speaker_name}」重新回复上一轮。要求：{payload.requirement or '严格符合人物设定和当前视角'}。输出合法 JSON messages 数组。"})
            params = generation_params(ctx.settings, api_config)
            handle.progress("角色单条消息重生成中", percent=35, stage="生成回复")
            response = client.chat.completions.create(model=model, messages=legacy, temperature=params["temperature"], top_p=params["top_p"], max_tokens=params["max_tokens"], frequency_penalty=params["frequency_penalty"])
            raw = response.choices[0].message.content or ""
            turn = max([int(getattr(item, "turn_index", 0) or 0) for item in branch_messages] or [0])
            generated = parse_structured_reply(raw, branch_data.get("branch_id") or "main", turn, {speaker_name: speaker_id})
            generated = [item for item in generated if item.speaker_id == speaker_id or item.speaker_name == speaker_name]
            if not generated:
                generated = [ChatMessage(message_id=new_id("msg"), branch_id=branch_data.get("branch_id") or "main", role="assistant", speaker_id=speaker_id, speaker_name=speaker_name, content=raw, turn_index=turn, created_at=now_text())]
            branch_data["messages"] = [*branch_data.get("messages", []), *[asdict(item) for item in generated]]
            updated = ctx.conversation_manager.load_conversation(conversation_id) or record
            updated = ensure_roleplay_branches(updated)
            for item in updated.get("branches") or []:
                if item.get("branch_id") == branch_data.get("branch_id"):
                    item.update(branch_data)
                    break
            updated["active_branch_id"] = branch_data.get("branch_id") or updated.get("active_branch_id") or "main"
            updated["structured_messages"] = list(branch_data.get("messages") or [])
            save_roleplay_record(ctx, updated)
            result = {"conversation": ctx.conversation_manager.load_conversation(conversation_id) or updated, "messages": branch_data.get("messages") or [], "assistant_messages": [asdict(item) for item in generated]}
            handle.progress("角色消息重生成完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"重生成角色消息《{conversation_id}》", target, metadata={"kind": "roleplay_message_regenerate", "conversation_id": conversation_id, "message_id": message_id}, retryable=True)}

    @app.post("/api/roleplay/conversations/{conversation_id}/export", tags=["roleplay"])
    def export_roleplay_conversation(conversation_id: str, payload: ConversationExportRequest, ctx: WebUserContext = Depends(current_context)):
        fmt = normalize_fmt(payload.fmt)
        runtime.cleanup_exports(ctx)
        record = ctx.conversation_manager.load_conversation(conversation_id)
        if not record:
            raise HTTPException(status_code=404, detail="会话不存在")
        out = os.path.join(ctx.export_root, f"{safe_name(record.get('title') or conversation_id)}_conversation.{fmt}")
        path = export_conversation(ctx.conversation_manager, conversation_id, fmt, out)
        download = runtime.register_download(ctx.username, path, os.path.basename(path), media_type_for(path))
        return {"download": download}

    @app.post("/api/roleplay/conversations", tags=["roleplay"])
    def save_conversation(payload: ConversationSaveRequest, ctx: WebUserContext = Depends(current_context)):
        record = dict(payload.record or {})
        conversation_id = record.get("conversation_id") or ctx.conversation_manager.generate_id(record.get("title") or "角色对话")
        record["conversation_id"] = conversation_id
        save_roleplay_record(ctx, record)
        return {"conversation_id": conversation_id, "conversation": ctx.conversation_manager.load_conversation(conversation_id) or record}

    @app.post("/api/roleplay/chat", tags=["roleplay"])
    def roleplay_chat(payload: RoleChatRequest, ctx: WebUserContext = Depends(current_context)):
        try:
            api_config, client, model = text_client_and_model(ctx)
        except WebApiConfigError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        def target(handle):
            from core.character_book import dict_to_timeline, timeline_to_dict
            from core.chat_domain import ChatMessage, legacy_messages_to_structured, parse_structured_reply, structured_to_legacy_messages, now_text, new_id
            from strategies.role_play_strategy import RolePlayStrategy
            book = ctx.character_book_manager.load()
            profiles = {profile.character_id: profile for profile in book.profiles}
            participant_ids = [cid for cid in payload.character_ids if cid in profiles]
            if not participant_ids and book.profiles:
                participant_ids = [book.profiles[0].character_id]
            if not participant_ids:
                raise RuntimeError("请先创建并选择至少一个角色")
            chat_type = "group" if payload.chat_type == "group" and len(participant_ids) > 1 else "private"
            conversation_id = payload.conversation_id.strip()
            record = ctx.conversation_manager.load_conversation(conversation_id) if conversation_id else None
            if not conversation_id:
                conversation_id = ctx.conversation_manager.generate_id(payload.title or "角色对话")
            name_to_id = {profile.name: profile.character_id for profile in book.profiles if profile.name}
            if record:
                record = ensure_roleplay_branches(record)
            scene_state = SceneState(**filter_fields(SceneState, payload.scene_state or (record or {}).get("scene_state") or {}))
            turn_policy = TurnPolicy(**filter_fields(TurnPolicy, payload.turn_policy or (record or {}).get("turn_policy") or {}))
            required_ids = [cid for cid in (payload.required_responder_ids or turn_policy.required_speaker_ids) if cid in participant_ids]
            if chat_type == "private":
                required_ids = participant_ids[:1]
            elif not required_ids:
                required_ids = list(participant_ids)
            turn_policy.required_speaker_ids = list(required_ids)
            turn_policy.allowed_speaker_ids = [cid for cid in turn_policy.allowed_speaker_ids if cid in participant_ids]
            turn_policy.blocked_speaker_ids = [cid for cid in turn_policy.blocked_speaker_ids if cid in participant_ids]
            turn_policy.speaker_order = [cid for cid in turn_policy.speaker_order if cid in participant_ids]
            turn_policy.mention_only_ids = [cid for cid in turn_policy.mention_only_ids if cid in participant_ids]
            scene_state.present_character_ids = [cid for cid in scene_state.present_character_ids if cid in participant_ids]
            sender_profile_id = payload.sender_profile_id or (record or {}).get("sender_profile_id") or ""
            sender_profile_record = next((item for item in ctx.sender_profile_manager.load() if item.sender_profile_id == sender_profile_id), None)
            sender_name = payload.sender_name or (record or {}).get("sender_name") or (sender_profile_record.name if sender_profile_record else "你")
            sender_profile_text = payload.sender_profile or (record or {}).get("sender_profile") or ""
            if sender_profile_record and not sender_profile_text.strip():
                sender_profile_text = "\n".join(str(value) for value in [sender_profile_record.identity, sender_profile_record.personality, sender_profile_record.appearance, sender_profile_record.background, sender_profile_record.relationships, sender_profile_record.knowledge_state, sender_profile_record.notes] if value)
            active_branch_id = (record or {}).get("active_branch_id") or "main"
            active_branch = next((item for item in (record or {}).get("branches", []) if item.get("branch_id") == active_branch_id), None)
            structured = []
            if active_branch and active_branch.get("messages"):
                structured = [ChatMessage(**filter_chat_message(item)) for item in active_branch.get("messages") or []]
            elif record and record.get("structured_messages"):
                structured = [ChatMessage(**filter_chat_message(item)) for item in record.get("structured_messages") or []]
            elif record:
                structured = legacy_messages_to_structured(record.get("messages") or [], branch_id=active_branch_id, sender_name=record.get("sender_name") or sender_name or "你", name_to_id=name_to_id)
            turn = max([int(getattr(item, "turn_index", 0) or 0) for item in structured] or [0]) + 1
            user_msg = ChatMessage(message_id=new_id("msg"), branch_id=active_branch_id, role="user", speaker_id="sender", speaker_name=sender_name or "你", content=payload.message, turn_index=turn, created_at=now_text())
            prompt_messages = [*structured, user_msg]
            strategy = RolePlayStrategy()
            strategy.character_book = book
            strategy.participant_character_ids = participant_ids
            strategy.primary_character_id = participant_ids[0]
            strategy.chat_type = chat_type
            strategy.required_responder_ids = required_ids
            strategy.sender_name = sender_name or "你"
            strategy.sender_profile = sender_profile_text or ""
            strategy.sender_profile_record = sender_profile_record
            strategy.scene_state = scene_state
            strategy.turn_policy = turn_policy
            strategy.reply_mode = payload.reply_mode if payload.reply_mode in {RolePlayStrategy.REPLY_MODE_CHARACTER, RolePlayStrategy.REPLY_MODE_NARRATOR} else RolePlayStrategy.REPLY_MODE_CHARACTER
            strategy.narrator_enabled = bool(payload.narrator_enabled)
            strategy.timeline = dict_to_timeline((record or {}).get("timeline") or [])
            messages = structured_to_legacy_messages(prompt_messages, strategy.get_system_prompt())
            params = generation_params(ctx.settings, api_config)
            params["model"] = model
            handle.progress("角色正在回复", percent=30, stage="生成回复")
            response = client.chat.completions.create(model=model, messages=messages, temperature=params["temperature"], top_p=params["top_p"], max_tokens=params["max_tokens"], frequency_penalty=params["frequency_penalty"])
            raw = response.choices[0].message.content or ""
            assistant_messages = parse_structured_reply(raw, active_branch_id, turn, name_to_id)
            if not assistant_messages:
                assistant_messages = [ChatMessage(message_id=new_id("msg"), branch_id=active_branch_id, role="assistant", speaker_id=participant_ids[0], speaker_name=profiles[participant_ids[0]].name, content=raw, turn_index=turn, created_at=now_text())]
            all_structured = [*structured, user_msg, *assistant_messages]
            branches = list((record or {}).get("branches") or [])
            if not branches:
                branches = [{"branch_id": active_branch_id, "title": "主线", "messages": [], "timeline": (record or {}).get("timeline") or [], "created_at": now_text()}]
            active_branch = next((item for item in branches if item.get("branch_id") == active_branch_id), None)
            if active_branch is None:
                active_branch = {"branch_id": active_branch_id, "title": "主线", "messages": [], "timeline": (record or {}).get("timeline") or [], "created_at": now_text()}
                branches.append(active_branch)
            active_branch["messages"] = [asdict(item) for item in all_structured]
            active_branch["timeline"] = timeline_to_dict(strategy.timeline)
            legacy = structured_to_legacy_messages(all_structured)
            title = payload.title or (record or {}).get("title") or "角色对话"
            ctx.conversation_manager.save_conversation(
                conversation_id=conversation_id,
                title=title,
                model=model,
                messages=legacy,
                strategy="角色扮演",
                reply_mode=strategy.reply_mode,
                chat_type=chat_type,
                participant_character_ids=participant_ids,
                primary_character_id=participant_ids[0],
                timeline_id=conversation_id,
                timeline=timeline_to_dict(strategy.timeline),
                character_book_snapshot=character_book_to_dict(book),
                sender_name=strategy.sender_name,
                sender_profile=strategy.sender_profile,
                required_responder_ids=required_ids,
                structured_messages=[asdict(item) for item in all_structured],
                branches=branches,
                active_branch_id=active_branch_id,
                sender_profile_id=sender_profile_id,
                scene_state=asdict(scene_state),
                turn_policy=asdict(turn_policy),
                memory_change_sets=(record or {}).get("memory_change_sets") or [],
                narrator_enabled=strategy.narrator_enabled,
            )
            result = {"conversation_id": conversation_id, "title": title, "messages": [asdict(item) for item in all_structured], "assistant_messages": [asdict(item) for item in assistant_messages]}
            handle.progress("角色回复完成", percent=100, stage="完成", data={"result": result})
            return result
        return {"task_id": runtime.start_task(ctx.username, f"角色聊天《{payload.title}》", target, metadata={"kind": "roleplay_chat", "conversation_id": payload.conversation_id}, retryable=True)}

    @app.get("/api/tasks", tags=["tasks"])
    def list_tasks(ctx: WebUserContext = Depends(current_context)):
        return {"tasks": runtime.list_tasks(ctx.username)}

    @app.get("/api/tasks/{task_id}", tags=["tasks"])
    def get_task(task_id: str, ctx: WebUserContext = Depends(current_context)):
        if not runtime.user_owns_task(ctx.username, task_id):
            raise HTTPException(status_code=404, detail="任务不存在")
        task = runtime.serialize_task(task_id)
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")
        return {"task": task}

    @app.post("/api/tasks/{task_id}/cancel", tags=["tasks"])
    def cancel_task(task_id: str, ctx: WebUserContext = Depends(current_context)):
        return {"ok": runtime.cancel_task(ctx.username, task_id)}

    @app.get("/api/tasks/{task_id}/events", tags=["tasks"])
    async def task_events(request: Request, task_id: str, token: str = Query(default="")):
        try:
            ctx = runtime.context_from_token(token)
        except WebAuthError as exc:
            raise HTTPException(status_code=401, detail=str(exc)) from exc
        if not runtime.user_owns_task(ctx.username, task_id):
            raise HTTPException(status_code=404, detail="任务不存在")
        async def event_stream():
            replayed_terminal = False
            for event in runtime.events_snapshot(task_id):
                yield task_event_to_sse(event)
                if event.type in {"completed", "failed", "cancelled"}:
                    replayed_terminal = True
            if replayed_terminal:
                return
            q = runtime.event_queue(task_id)
            while not await request.is_disconnected():
                try:
                    event = await asyncio.to_thread(q.get, True, 15)
                except queue.Empty:
                    yield ": keepalive\n\n"
                    continue
                yield task_event_to_sse(event)
                if event.type in {"completed", "failed", "cancelled"}:
                    break
        return StreamingResponse(event_stream(), media_type="text/event-stream")

    @app.post("/api/tasks/{task_id}/retry", tags=["tasks"])
    def retry_task(task_id: str, ctx: WebUserContext = Depends(current_context)):
        try:
            new_task_id = runtime.retry_task(ctx.username, task_id)
            return {"task_id": new_task_id}
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
    @app.get("/api/downloads/{download_id}", tags=["downloads"])
    def download(download_id: str, token: str = Query(default="")):
        try:
            ctx = runtime.context_from_token(token)
            item = runtime.resolve_download(ctx.username, download_id)
        except (WebAuthError, FileNotFoundError) as exc:
            raise HTTPException(status_code=404, detail="下载不存在或已过期") from exc
        return FileResponse(item["path"], media_type=item.get("media_type"), filename=item.get("filename"))

    @app.get("/api/token-log", tags=["diagnostics"])
    def token_log(q: str = "", model: str = "", operation: str = "", date_from: str = "", date_to: str = "", ctx: WebUserContext = Depends(current_context)):
        all_entries = [asdict(item) for item in ctx.token_log_manager.list_entries()]
        entries = filter_token_entries(all_entries, q=q, model=model, operation=operation, date_from=date_from, date_to=date_to)
        summary = token_summary(entries)
        return {"entries": entries, "total": len(entries), "overall_total": len(all_entries), "summary": summary, "facets": token_facets(all_entries), "filters": {"q": q, "model": model, "operation": operation, "date_from": date_from, "date_to": date_to}}

    @app.delete("/api/token-log", tags=["diagnostics"])
    def clear_token_log(ctx: WebUserContext = Depends(current_context)):
        ctx.token_log_manager.clear()
        return {"ok": True}

    @app.post("/api/token-log/export", tags=["diagnostics"])
    def export_token_log(q: str = "", model: str = "", operation: str = "", date_from: str = "", date_to: str = "", ctx: WebUserContext = Depends(current_context)):
        runtime.cleanup_exports(ctx)
        all_entries = [asdict(item) for item in ctx.token_log_manager.list_entries()]
        entries = filter_token_entries(all_entries, q=q, model=model, operation=operation, date_from=date_from, date_to=date_to)
        path = os.path.join(ctx.export_root, "token_log.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"entries": entries, "summary": token_summary(entries), "filters": {"q": q, "model": model, "operation": operation, "date_from": date_from, "date_to": date_to}}, f, ensure_ascii=False, indent=2)
        download = runtime.register_download(ctx.username, path, "token_log.json", "application/json")
        return {"download": download}

    @app.get("/api/diagnostics", tags=["diagnostics"])
    def diagnostics(ctx: WebUserContext = Depends(current_context)):
        return build_diagnostics_payload(runtime, ctx)

    @app.post("/api/diagnostics/export", tags=["diagnostics"])
    def export_diagnostics(ctx: WebUserContext = Depends(current_context)):
        runtime.cleanup_exports(ctx)
        payload = build_diagnostics_payload(runtime, ctx)
        path = os.path.join(ctx.export_root, "diagnostics.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        download = runtime.register_download(ctx.username, path, "diagnostics.json", "application/json")
        return {"download": download}

    static_dir = Path(__file__).resolve().parent / "static"
    app.mount("/", StaticFiles(directory=static_dir, html=True), name="static")
    return app


def clamp_int(value, default: int, low: int, high: int) -> int:
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = default
    return max(low, min(high, number))


def normalize_preset(data: dict) -> dict:
    data = data or {}
    return {
        "temp": clamp_int(data.get("temp"), 70, 0, 200),
        "top_p": clamp_int(data.get("top_p"), 90, 0, 100),
        "fp": clamp_int(data.get("fp"), 0, -200, 200),
        "max_tokens": clamp_int(data.get("max_tokens"), 32768, 1, 300000),
    }

def password_strength_ok(password: str) -> bool:
    return len(password or "") >= 6 and any(ch.isalpha() for ch in password) and any(ch.isdigit() for ch in password)


async def read_single_upload(request: Request) -> dict:
    content_type = request.headers.get("content-type", "")
    body = await request.body()
    if not body:
        return {}
    if "multipart/form-data" not in content_type.lower():
        return {"filename": "upload.zip", "content": body}
    message = BytesParser(policy=email_policy).parsebytes((f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n").encode("utf-8") + body)
    for part in message.iter_parts():
        filename = part.get_filename()
        if filename:
            return {"filename": filename, "content": part.get_payload(decode=True) or b""}
    return {}

def serialize_meta(meta: NovelMeta) -> dict:
    return {"title": meta.title, "author": meta.author, "protagonist_bio": meta.protagonist_bio, "background_story": meta.background_story, "writing_demand": meta.writing_demand, "author_plan": meta.author_plan, "genre": meta.genre, "style_tone": meta.style_tone, "xp_mode": bool(meta.xp_mode), "total_chapters": meta.total_chapters, "created_at": meta.created_at, "updated_at": meta.updated_at}


def normalize_chapter(item: dict) -> dict:
    data = dict(item)
    num = data.get("chapter_num", data.get("num", 0))
    data["chapter_num"] = int(num or 0)
    data.setdefault("num", data["chapter_num"])
    data.setdefault("title", data.get("chapter_title") or f"第{data['chapter_num']}章")
    return data


def model_data(model: BaseModel) -> dict:
    if hasattr(model, "model_dump"):
        return model.model_dump()
    return model.dict()


def active_roleplay_branch(record: dict) -> dict:
    record = ensure_roleplay_branches(record)
    active_id = record.get("active_branch_id") or "main"
    branch = next((item for item in record.get("branches", []) if item.get("branch_id") == active_id), None)
    if not branch:
        raise HTTPException(status_code=404, detail="活跃分支不存在")
    return branch

def ensure_roleplay_branches(record: dict) -> dict:
    record = dict(record or {})
    branches = list(record.get("branches") or [])
    active_id = record.get("active_branch_id") or "main"
    if not branches:
        branches = [{
            "branch_id": active_id,
            "title": "主线",
            "parent_branch_id": "",
            "fork_message_id": "",
            "messages": list(record.get("structured_messages") or []),
            "timeline": list(record.get("timeline") or []),
            "character_state_snapshot": {},
            "knowledge": [],
            "relationships": [],
            "created_at": str(record.get("created_at") or ""),
        }]
    if not any(item.get("branch_id") == active_id for item in branches):
        active_id = branches[0].get("branch_id") or "main"
    for branch in branches:
        branch.setdefault("messages", [])
        branch.setdefault("timeline", [])
        branch.setdefault("title", branch.get("branch_id") or "分支")
    record["branches"] = branches
    record["active_branch_id"] = active_id
    active = next((item for item in branches if item.get("branch_id") == active_id), branches[0])
    record["structured_messages"] = list(active.get("messages") or record.get("structured_messages") or [])
    return record


def save_roleplay_record(ctx: WebUserContext, record: dict) -> None:
    record = ensure_roleplay_branches(record)
    ctx.conversation_manager.save_conversation(
        conversation_id=record.get("conversation_id") or ctx.conversation_manager.generate_id(record.get("title") or "角色对话"),
        title=record.get("title") or "角色对话",
        model=record.get("model") or "",
        messages=record.get("messages") or [],
        strategy=record.get("strategy") or "角色扮演",
        reply_mode=record.get("reply_mode") or "character",
        chat_type=record.get("chat_type") or "private",
        participant_character_ids=record.get("participant_character_ids") or [],
        primary_character_id=record.get("primary_character_id") or "",
        timeline_id=record.get("timeline_id") or record.get("conversation_id") or "",
        timeline=record.get("timeline") or [],
        character_book_snapshot=record.get("character_book_snapshot") or {},
        sender_name=record.get("sender_name") or "你",
        sender_profile=record.get("sender_profile") or "",
        required_responder_ids=record.get("required_responder_ids") or [],
        structured_messages=record.get("structured_messages") or [],
        branches=record.get("branches") or [],
        active_branch_id=record.get("active_branch_id") or "main",
        sender_profile_id=record.get("sender_profile_id") or "",
        scene_state=record.get("scene_state") or {},
        turn_policy=record.get("turn_policy") or {},
        memory_change_sets=record.get("memory_change_sets") or [],
        narrator_enabled=bool(record.get("narrator_enabled", False)),
    )

def find_memory_change_set(record: dict, change_set_id: str):
    for index, item in enumerate(record.get("memory_change_sets") or []):
        if item.get("change_set_id") == change_set_id:
            return index, change_set_from_dict(item)
    raise HTTPException(status_code=404, detail="记忆变更不存在")


def update_memory_change_set(record: dict, index: int, change_set) -> None:
    changes = list(record.get("memory_change_sets") or [])
    while len(changes) <= index:
        changes.append({})
    changes[index] = asdict(change_set)
    record["memory_change_sets"] = changes


def snapshot_roleplay_character_book(record: dict, book) -> None:
    snapshot = character_book_to_dict(book)
    record["character_book_snapshot"] = snapshot
    record = ensure_roleplay_branches(record)
    active_id = record.get("active_branch_id") or "main"
    for branch in record.get("branches") or []:
        if branch.get("branch_id") == active_id:
            branch["character_state_snapshot"] = snapshot
            branch["timeline"] = list(record.get("timeline") or branch.get("timeline") or [])
            break


def roleplay_memory_state(ctx: WebUserContext, record: dict) -> dict:
    record = ensure_roleplay_branches(record)
    return {
        "conversation_id": record.get("conversation_id", ""),
        "active_branch_id": record.get("active_branch_id", "main"),
        "timeline": record.get("timeline") or [],
        "memory_change_sets": record.get("memory_change_sets") or [],
        "character_book_snapshot": record.get("character_book_snapshot") or {},
        "book": character_book_to_dict(ctx.character_book_manager.load()),
    }
def roleplay_control_state(record: dict) -> dict:
    return {
        "conversation_id": record.get("conversation_id", ""),
        "chat_type": record.get("chat_type", "private"),
        "reply_mode": record.get("reply_mode", "character"),
        "participant_character_ids": record.get("participant_character_ids") or [],
        "required_responder_ids": record.get("required_responder_ids") or [],
        "sender_name": record.get("sender_name", ""),
        "sender_profile": record.get("sender_profile", ""),
        "sender_profile_id": record.get("sender_profile_id", ""),
        "scene_state": record.get("scene_state") or {},
        "turn_policy": record.get("turn_policy") or {},
        "active_branch_id": record.get("active_branch_id", "main"),
        "narrator_enabled": bool(record.get("narrator_enabled", False)),
        "timeline": record.get("timeline") or [],
        "memory_change_sets": record.get("memory_change_sets") or [],
    }


def filter_profile(data: dict) -> dict:
    return {key: value for key, value in (data or {}).items() if key in CharacterProfile.__dataclass_fields__}


def safe_name(text: str) -> str:
    return "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in text)[:80] or "export"


def normalize_fmt(fmt: str) -> str:
    fmt = (fmt or "txt").lower().strip(".")
    if fmt not in {"txt", "md", "html", "docx"}:
        raise HTTPException(status_code=400, detail="不支持的导出格式")
    return fmt


def normalize_context_policy(policy: dict) -> dict:
    policy = dict(policy or {})
    mode = str(policy.get("load_mode") or "auto")
    if mode not in {"resident", "auto", "manual"}:
        mode = "auto"
    try:
        priority = int(policy.get("priority", 50))
    except (TypeError, ValueError):
        priority = 50
    keywords = policy.get("keywords") or []
    if isinstance(keywords, str):
        keywords = [item.strip() for item in keywords.replace(",", "、").split("、") if item.strip()]
    elif isinstance(keywords, list):
        keywords = [str(item).strip() for item in keywords if str(item).strip()]
    else:
        keywords = []
    return {
        "enabled": bool(policy.get("enabled", True)),
        "load_mode": mode,
        "priority": max(0, min(100, priority)),
        "brief_description": str(policy.get("brief_description") or "").strip(),
        "keywords": keywords,
    }



def build_diagnostics_payload(runtime: WebRuntime, ctx: WebUserContext) -> dict:
    token_entries = [asdict(item) for item in ctx.token_log_manager.list_entries()]
    tasks = runtime.list_tasks(ctx.username, limit=100)
    settings_keys = sorted(ctx.settings.keys())
    api_config = ctx.load_api_config()
    text_api = api_config.get("text") or {}
    downloads = sorted(
        [item for item in runtime._downloads.values() if item.get("username") == ctx.username],
        key=lambda item: str(item.get("created_at") or ""),
        reverse=True,
    )[:30]
    return {
        "schema_version": 1,
        "user": ctx.username,
        "time": time.strftime("%Y-%m-%d %H:%M:%S"),
        "books": len(ctx.novel_manager.list_books()),
        "conversations": len(ctx.conversation_manager.list_conversations()),
        "api_configured": bool(text_api.get("api_key")),
        "text_model": text_api.get("model") or ctx.settings.get("last_model") or "",
        "settings_keys": settings_keys,
        "tasks": tasks,
        "task_summary": {
            "total": len(tasks),
            "running": len([item for item in tasks if item.get("status") == "running"]),
            "failed": len([item for item in tasks if item.get("status") == "failed"]),
            "completed": len([item for item in tasks if item.get("status") == "completed"]),
        },
        "token_summary": token_summary(token_entries),
        "token_log_count": len(token_entries),
        "downloads": [
            {
                "download_id": item.get("download_id"),
                "download_url": item.get("download_url"),
                "filename": item.get("filename"),
                "created_at": item.get("created_at"),
            }
            for item in downloads
        ],
    }

def filter_token_entries(entries: list[dict], *, q: str = "", model: str = "", operation: str = "", date_from: str = "", date_to: str = "") -> list[dict]:
    keyword = (q or "").strip().lower()
    model_filter = (model or "").strip()
    operation_filter = (operation or "").strip()
    start = (date_from or "").strip()
    end = (date_to or "").strip()
    result = []
    for row in entries:
        stamp = str(row.get("timestamp") or "")
        day = stamp[:10]
        if model_filter and str(row.get("model") or "") != model_filter:
            continue
        if operation_filter and str(row.get("operation") or "") != operation_filter:
            continue
        if start and day < start:
            continue
        if end and day > end:
            continue
        if keyword:
            haystack = " ".join(str(row.get(key) or "") for key in ("timestamp", "operation", "direction", "strategy", "model", "content_preview", "usage_status")).lower()
            if keyword not in haystack:
                continue
        result.append(row)
    return result


def token_facets(entries: list[dict]) -> dict:
    models = sorted({str(row.get("model") or "未指定") for row in entries})
    operations = sorted({str(row.get("operation") or "未指定") for row in entries})
    return {"models": models, "operations": operations}
def estimate_token_cost(prompt_tokens: int, completion_tokens: int, model: str = "") -> dict:
    return {
        "currency": "USD",
        "input_cost": 0.0,
        "output_cost": 0.0,
        "total_cost": 0.0,
        "pricing": "not_configured",
        "model": model or "mixed",
    }


def token_summary(entries: list[dict]) -> dict:
    totals = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    activity = {"duration_ms": 0, "duration_count": 0, "char_count": 0, "char_count_entries": 0, "hanzi_count": 0, "hanzi_count_entries": 0}
    by_model: dict[str, dict] = {}
    by_operation: dict[str, dict] = {}
    by_date: dict[str, dict] = {}

    def add_activity(row: dict) -> None:
        if row.get("duration_ms") is not None:
            activity["duration_ms"] += int(row.get("duration_ms") or 0)
            activity["duration_count"] += 1
        if row.get("char_count") is not None:
            activity["char_count"] += int(row.get("char_count") or 0)
            activity["char_count_entries"] += 1
        if row.get("hanzi_count") is not None:
            activity["hanzi_count"] += int(row.get("hanzi_count") or 0)
            activity["hanzi_count_entries"] += 1

    def add(bucket: dict, key: str, row: dict) -> None:
        item = bucket.setdefault(key or "未分类", {"count": 0, "prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "duration_ms": 0, "char_count": 0, "hanzi_count": 0})
        item["count"] += 1
        for name in totals:
            item[name] += int(row.get(name) or 0)
        for name in ("duration_ms", "char_count", "hanzi_count"):
            item[name] += int(row.get(name) or 0)

    for row in entries:
        if row.get("usage_status") == "ok":
            for name in totals:
                totals[name] += int(row.get(name) or 0)
        add_activity(row)
        add(by_model, str(row.get("model") or "未指定"), row)
        add(by_operation, str(row.get("operation") or "未指定"), row)
        add(by_date, str(row.get("timestamp") or "")[:10], row)

    for model_name, item in by_model.items():
        item["estimated_cost"] = estimate_token_cost(item.get("prompt_tokens", 0), item.get("completion_tokens", 0), model_name)
    for bucket in (by_operation, by_date):
        for item in bucket.values():
            item["estimated_cost"] = estimate_token_cost(item.get("prompt_tokens", 0), item.get("completion_tokens", 0))
    return {
        "totals": totals,
        "activity": activity,
        "estimated_cost": estimate_token_cost(totals["prompt_tokens"], totals["completion_tokens"]),
        "by_model": by_model,
        "by_operation": by_operation,
        "by_date": by_date,
    }

def section_dicts(sections) -> list[dict]:
    return [{"title": str(title or f"分段 {idx}").strip(), "content": str(content or "").strip()} for idx, (title, content) in enumerate(sections or [], 1) if str(content or "").strip()]


def normalize_sections(sections: list[dict]) -> list[dict]:
    result = []
    for idx, section in enumerate(sections or [], 1):
        title = str(section.get("title") or f"分段 {idx}").strip()
        content = str(section.get("content") or "").strip()
        if content:
            result.append({"title": title, "content": content})
    return result


def continuation_upload_root(ctx: WebUserContext) -> str:
    root = os.path.join(ctx.user_dir, ".deepseekass", "web_uploads", "continuation")
    os.makedirs(root, exist_ok=True)
    return root


def cleanup_old_files(root: str, max_age_seconds: int = 24 * 60 * 60) -> None:
    cutoff = time.time() - max_age_seconds
    for name in os.listdir(root):
        path = os.path.join(root, name)
        try:
            if os.path.isfile(path) and os.path.getmtime(path) < cutoff:
                os.remove(path)
        except OSError:
            pass


def safe_upload_filename(filename: str) -> str:
    name = os.path.basename((filename or "source.txt").replace("\\", "/"))
    stem, ext = os.path.splitext(name)
    ext = ext.lower() if ext.lower() in {".txt", ".md", ".markdown", ".html", ".htm"} else ".txt"
    return f"{safe_name(stem)}{ext}"


def decode_uploaded_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def sort_upload_items(items: list[dict]) -> list[dict]:
    def key(item):
        name = str(item.get("filename") or "")
        nums = re.findall(r"\d+", name)
        return (int(nums[0]) if nums else 10**9, name.lower())
    return sorted(items, key=key)


async def read_multipart_text_files(request: Request) -> list[dict]:
    content_type = request.headers.get("content-type", "")
    body = await request.body()
    if not body:
        return []
    if "multipart/form-data" not in content_type.lower():
        return [{"filename": "pasted.txt", "content": decode_uploaded_text(body)}]
    message = BytesParser(policy=email_policy).parsebytes((f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n").encode("utf-8") + body)
    result = []
    for part in message.iter_parts():
        filename = part.get_filename()
        if not filename:
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in {".txt", ".md", ".markdown", ".html", ".htm"}:
            continue
        result.append({"filename": filename, "content": decode_uploaded_text(part.get_payload(decode=True) or b"")})
    return result


def world_counts(world_data: dict) -> dict:
    keys = ["characters", "locations", "organizations", "items", "rules", "world_rules", "timeline", "plot_threads", "active_plot_threads", "global_foreshadowing"]
    return {key: len(world_data.get(key) or []) for key in keys if isinstance(world_data.get(key), list)}


def build_continuation_prompt(*, title: str, chapter_num: int, chapter_title: str, source_text: str, requirement: str, plot: str, setting: str, target_words: int, meta: NovelMeta, context_text: str, xp_mode: bool = False, chapter_mode: bool = True) -> str:
    if chapter_mode:
        parts = [
            f"请基于旧文和当前书籍资料，为小说《{title}》续写第{chapter_num}章《{chapter_title}》。",
            f"目标字数：约 {target_words} 字。正文必须自然承接前文，保留人物关系、语气、伏笔和世界规则。",
        ]
    else:
        parts = [
            f"请基于旧文和当前书籍资料，为小说《{title}》生成一份续写草稿《{chapter_title}》。",
            f"目标字数：约 {target_words} 字。正文必须自然承接前文，保留人物关系、语气、伏笔和世界规则。不要写入章节树，也不要假定这是正式章节版本。",
        ]
    if setting.strip():
        parts.append(f"用户补充设定：\n{setting.strip()}")
    if requirement.strip():
        parts.append(f"续写要求：\n{requirement.strip()}")
    if plot.strip():
        parts.append(f"指定剧情/发展方向：\n{plot.strip()}")
    meta_parts = [meta.background_story, meta.protagonist_bio, meta.writing_demand, meta.author_plan]
    meta_text = "\n\n".join([part for part in meta_parts if isinstance(part, str) and part.strip()])
    if meta_text:
        parts.append(f"当前小说设定：\n{meta_text}")
    parts.append(f"桌面端同源上下文：\n{context_text}")
    if source_text.strip():
        parts.append(f"源文档末尾参考：\n{source_text[-6000:]}")
    if xp_mode:
        parts.append("成人向 XP 模式已开启：在合法合规前提下保留用户设定的题材尺度、氛围和人物欲望驱动。")
    parts.append("只输出续写正文，不要输出解释、提纲、Markdown 标题或附加说明。")
    return "\n\n".join(parts)

def save_continuation_run(ctx: WebUserContext, book_title: str, task: str, *, input_chars: int = 0, input_summary: dict | None = None, output_summary: dict | None = None, result: dict | None = None, status: str = "completed", selected_skills: list | None = None) -> dict:
    from core.agent.types import now_iso
    title = (book_title or "").strip()
    if not title:
        return {}
    workspace = ctx.novel_manager.get_workspace(title)
    run_id = f"cont_web_{uuid.uuid4().hex}"
    record = {
        "schema_version": 1,
        "run_id": run_id,
        "task": task,
        "book_title": title,
        "selected_skills": selected_skills or [],
        "input_chars": int(input_chars or 0),
        "input_summary": input_summary or {},
        "output_summary": output_summary or {},
        "result": result or {},
        "status": status,
        "created_at": now_iso(),
    }
    workspace.storage.write_json(f"{workspace.agent_root}/continuation_runs/{run_id}.json", record)
    return record


def list_continuation_runs(ctx: WebUserContext, title: str = "") -> list[dict]:
    books = [title] if title else ctx.novel_manager.list_books()
    runs = []
    for book in books:
        if not book:
            continue
        try:
            workspace = ctx.novel_manager.get_workspace(book)
            for rel in workspace.storage.list_files(f"{workspace.agent_root}/continuation_runs"):
                if not str(rel).endswith(".json"):
                    continue
                data = workspace.storage.read_json(rel, default={}) or {}
                data.setdefault("book_title", book)
                data.setdefault("schema_version", 1)
                runs.append(data)
        except Exception:
            continue
    runs.sort(key=lambda item: str(item.get("created_at") or ""), reverse=True)
    return runs[:100]


def load_continuation_run(ctx: WebUserContext, run_id: str, *, title: str = "") -> dict | None:
    run_name = safe_name(os.path.basename(str(run_id or "")))
    if not run_name:
        return None
    books = [title] if title else ctx.novel_manager.list_books()
    for book in books:
        if not book:
            continue
        try:
            workspace = ctx.novel_manager.get_workspace(book)
            data = workspace.storage.read_json(f"{workspace.agent_root}/continuation_runs/{run_name}.json", default=None)
            if data:
                data.setdefault("book_title", book)
                data.setdefault("schema_version", 1)
                return data
        except Exception:
            continue
    return None

def filter_chat_message(data: dict) -> dict:
    from core.chat_domain import ChatMessage
    return {key: value for key, value in (data or {}).items() if key in ChatMessage.__dataclass_fields__}


def markdown_display_path(ctx: WebUserContext, full_path: str) -> str:
    root = os.path.abspath(ctx.markdown_root)
    full = os.path.abspath(full_path)
    rel = os.path.relpath(full, root).replace("\\", "/")
    if rel.endswith(".md.enc") or rel.endswith(".markdown.enc"):
        rel = rel[:-4]
    return rel


def markdown_is_note_file(ctx: WebUserContext, full_path: str) -> bool:
    name = os.path.basename(full_path).lower()
    if ctx.enc_key and name.endswith((".md.enc", ".markdown.enc")):
        return True
    return name.endswith((".md", ".markdown"))


def markdown_tree_items(ctx: WebUserContext) -> list[dict]:
    items = []
    root = os.path.abspath(ctx.markdown_root)
    os.makedirs(root, exist_ok=True)
    for current, dirs, files in os.walk(root):
        dirs[:] = sorted(dirs, key=lambda value: value.lower())
        rel_dir = os.path.relpath(current, root).replace("\\", "/")
        rel_dir = "" if rel_dir == "." else rel_dir
        for name in dirs:
            path = f"{rel_dir}/{name}".strip("/")
            items.append({"path": path, "name": name, "type": "folder"})
        for name in sorted(files, key=lambda value: value.lower()):
            full = os.path.join(current, name)
            if not markdown_is_note_file(ctx, full):
                continue
            display = markdown_display_path(ctx, full)
            items.append({"path": display, "name": os.path.basename(display), "type": "file"})
    return items


def media_type_for(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return {".txt": "text/plain; charset=utf-8", ".md": "text/markdown; charset=utf-8", ".html": "text/html; charset=utf-8", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".zip": "application/zip"}.get(ext, "application/octet-stream")


def write_node_export(path: str, fmt: str, title: str, node: dict, content: str) -> None:
    label = str(node.get("display_label") or node.get("title") or node.get("id") or "chapter")
    heading = f"{title} - {label}"
    nl = chr(10)
    if fmt == "txt":
        Path(path).write_text(heading + nl + ("=" * 40) + nl + nl + content + nl, encoding="utf-8")
        return
    if fmt == "md":
        Path(path).write_text("# " + heading + nl + nl + content + nl, encoding="utf-8")
        return
    if fmt == "html":
        body = html.escape(content).replace(nl, "<br>" + nl)
        html_text = '<!doctype html><html><head><meta charset="utf-8"><title>' + html.escape(heading) + '</title></head><body><h1>' + html.escape(heading) + '</h1><main>' + body + '</main></body></html>'
        Path(path).write_text(html_text, encoding="utf-8")
        return
    if fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(heading, level=1)
        for para in content.split(nl):
            doc.add_paragraph(para)
        doc.save(path)
        return
    raise ValueError(f"unsupported export format: {fmt}")


app = create_app()
